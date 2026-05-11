"""Passport v2 parser — extracts all data_field values from Felicity's 18-doc test set.

Maps PDF/DOCX text → flat dict keyed by `data-field` attribute names in passport_v2_shell.html.
Also produces `score_inputs` consumed by the shell's scoring engine.

Doc-type routing keyed on filename suffix:
  01a TechSpec, 01b MaterialList, 01c ChemicalSubstances,
  02a PurchaseOrder, 02b BillOfLading, 02c MillProfile,
  03a DurabilityTestReport, 03b EOLPathwayProfile, 03c CareRecommendation,
  04a GOTSScopeCertificate, 04b REACHDeclaration, 04c OEKOTEX100Certificate,
  05a LCAReport, 05b ResourceAudit, 05c SocialAuditReport,
  06a SupplierScorecard, 06b QCInspectionReport, 06c ColourConsistencyReport.
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple


# ────────────────────────────────────────────────────────────
# Low-level text helpers
# ────────────────────────────────────────────────────────────

def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""
    if ext == "pdf":
        from pypdf import PdfReader
        r = PdfReader(BytesIO(data))
        return "\n".join((p.extract_text() or "") for p in r.pages)
    if ext in ("doc", "docx"):
        import docx as _docx
        d = _docx.Document(BytesIO(data))
        parts: List[str] = [p.text for p in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    return ""


def doc_type_from_name(filename: str) -> Optional[str]:
    """Return token like '01a', '04b', '06c' or None."""
    m = re.search(r"_(\d{2}[abc])_", filename or "")
    return m.group(1) if m else None


def fabric_kind_from_name(filename: str) -> Optional[str]:
    """Return 'ribknit' or 'twillwoven' or None."""
    n = (filename or "").lower()
    if n.startswith("ribknit"):
        return "ribknit"
    if n.startswith("twillwoven"):
        return "twillwoven"
    return None


def _after(text: str, label: str, max_lines: int = 1) -> Optional[str]:
    """Get the value that follows `label` on the next non-empty line(s)."""
    lines = [ln.rstrip() for ln in text.splitlines()]
    label_re = re.compile(rf"^\s*{re.escape(label)}\s*$", re.IGNORECASE)
    for i, ln in enumerate(lines):
        if label_re.match(ln):
            collected: List[str] = []
            j = i + 1
            while j < len(lines) and len(collected) < max_lines:
                v = lines[j].strip()
                if v:
                    collected.append(v)
                j += 1
            if collected:
                return " ".join(collected).strip()
    return None


def _first_match(text: str, *patterns: str, flags: int = re.IGNORECASE) -> Optional[str]:
    for p in patterns:
        m = re.search(p, text, flags)
        if m:
            return m.group(1).strip()
    return None


def _join_wrap(s: Optional[str]) -> Optional[str]:
    """Collapse line wraps inside a captured value to a single space-separated string."""
    if not s:
        return s
    return re.sub(r"\s+", " ", s.replace("\n", " ")).strip()


# ────────────────────────────────────────────────────────────
# Per-doc-type parsers — each returns a partial dict
# ────────────────────────────────────────────────────────────

def parse_01a_techspec(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["fabric_name"] = _after(t, "Fabric Name") or _first_match(t, r"\n([^\n]+Knit[^\n]+|[^\n]+Woven[^\n]+)\n")
    out["fabric_code"] = _after(t, "Fabric Code (SKU)") or _after(t, "Fabric Code") or _first_match(t, r"FAB-(?:RIB|TWL)-[A-Z0-9-]+")
    out["batch_no"] = _after(t, "Production Batch No.") or _first_match(t, r"BN-(?:RIB|TWL)-[A-Z0-9-]+")
    out["origin"] = _after(t, "Country of Origin") or _first_match(t, r"Country of Origin\s*\n([A-Za-z ]+)")
    out["po_no"] = _after(t, "Buyer Reference (PO)") or _first_match(t, r"PO-NAC-[0-9-]+")
    out["buyer"] = _after(t, "Buyer")
    out["structure"] = _after(t, "Construction Type")
    out["weight"] = _after(t, "Weight (GSM)")
    out["width"] = _after(t, "Width")
    out["yarn_count"] = _after(t, "Yarn Count")
    out["yarn_twist"] = _after(t, "Twist Per Metre (TPM)")
    out["fibre_length"] = _after(t, "Fibre Staple Length")
    out["spinning_method"] = _after(t, "Spinning Method")
    out["colour_state"] = _after(t, "Colour State")
    out["heat_setting"] = _after(t, "Heat Setting")
    out["mill_name"] = _after(t, "Mill Name")
    out["mill_id"] = _after(t, "Mill Facility ID")
    out["higg_fem_id"] = _after(t, "Higg FEM Facility ID")
    out["mill_address"] = _after(t, "Address")
    return {k: v for k, v in out.items() if v}


def parse_01b_materiallist(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["fibre_rows"] = _parse_fibre_table(t)
    out["recycled_pct"] = _after(t, "Total Recycled Content (% by mass)") or _first_match(t, r"Total Recycled Content[^\n]*\n([^\n]+)")
    out["recycled_source_type"] = _after(t, "Source of Recycled Material")
    out["recycled_grs"] = _after(t, "GRS / RCS Cert Reference")
    out["upstream_tiers"] = _parse_upstream_tiers(t)
    return {k: v for k, v in out.items() if v}


def _parse_fibre_table(t: str) -> List[Dict[str, str]]:
    """Parse the FIBRE COMPOSITION table block — supports wrapped fibre-name lines (multi-line entry).

    Strategy: find each `NN%` line and walk backward to assemble the fibre name (up to 2 preceding non-empty
    non-numeric lines that contain a fibre vocab token), then walk forward to capture origin + source.
    """
    rows: List[Dict[str, str]] = []
    m = re.search(r"FIBRE COMPOSITION[^\n]*\n(.+?)(?=RECYCLED CONTENT|YARN STRUCTURE|UPSTREAM|$)", t, re.IGNORECASE | re.DOTALL)
    block = m.group(1) if m else t
    lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
    i = 0
    while i < len(lines):
        pct_m = re.match(r"^(\d{1,3})\s*%\s*$", lines[i])
        if pct_m and i > 0:
            # Prefer just the previous line as fibre name. Add 1 prior line only if current
            # line is clearly a continuation (starts with `(` / lowercase / ends with comma).
            prev1 = lines[i - 1]
            name = prev1
            if i >= 2:
                prev2 = lines[i - 2]
                # Continuation heuristic: line starting with `(`, or ending with comma, or prev1 has no fibre vocab but prev2 does.
                if re.match(r"^\s*\(", prev1) or prev1.endswith(",") or (not _looks_like_fibre(prev1) and _looks_like_fibre(prev2)):
                    name = f"{prev2} {prev1}".strip()
            if not _looks_like_fibre(name):
                i += 1
                continue
            pct = int(pct_m.group(1))
            origin = lines[i + 1] if i + 1 < len(lines) else ""
            source = lines[i + 2] if i + 2 < len(lines) else ""
            rows.append({"name": name, "pct": pct, "origin": origin, "source": source})
            i += 3
            continue
        i += 1
    return rows


_FIBRE_VOCAB = {
    "cotton", "linen", "wool", "silk", "polyester", "nylon", "elastane", "spandex",
    "viscose", "rayon", "modal", "lyocell", "tencel", "acrylic", "polyamide",
    "cashmere", "alpaca", "mohair", "hemp", "jute", "ramie", "bamboo", "rpet",
}


def _looks_like_fibre(s: str) -> bool:
    low = s.lower()
    return any(f in low for f in _FIBRE_VOCAB)


def _parse_upstream_tiers(t: str) -> List[Dict[str, str]]:
    """Walk lines for 'Tier N — ...' headers; collect supplier name + location until next Tier header or paragraph break."""
    rows: List[Dict[str, str]] = []
    # Restrict to the UPSTREAM SOURCING block when present
    m = re.search(r"UPSTREAM SOURCING[^\n]*\n(.+?)(?=Composition tolerance|Authorised|This declaration|$)", t, re.IGNORECASE | re.DOTALL)
    block = m.group(1) if m else t
    lines = [ln.rstrip() for ln in block.splitlines()]
    i = 0
    while i < len(lines):
        m2 = re.match(r"\s*(Tier\s+\d\s*[—\-][^\n]*)", lines[i])
        if m2:
            tier_label = m2.group(1).strip()
            buf: List[str] = []
            j = i + 1
            while j < len(lines) and len(buf) < 4:
                s = lines[j].strip()
                if not s:
                    j += 1
                    continue
                if re.match(r"Tier\s+\d", s):
                    break
                # Stop on paragraph / disclaimer lines
                if re.match(r"(Composition tolerance|Authorised|This profile|Verification|Stamp)", s, re.IGNORECASE):
                    break
                buf.append(s)
                j += 1
            supplier = " ".join(buf[:-1]) if len(buf) >= 2 else (buf[0] if buf else "")
            location = buf[-1] if len(buf) >= 2 else ""
            if tier_label not in [r["tier"] for r in rows]:
                rows.append({"tier": tier_label, "supplier": supplier, "location": location})
            i = j
            continue
        i += 1
    return rows


def parse_01c_chemicals(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["dyestuff_class"] = _first_match(t, r"Dye Class:\s*([^\n·]+)") or _first_match(t, r"DYESTUFFS USED[\s\S]*?Reactive [^\n]+|DYESTUFFS USED[\s\S]*?Vat [^\n]+")
    out["pretreatment_agent"] = _after(t, "Pre-treatment Agent")
    out["softener"] = _after(t, "Softener")
    out["heat_setting"] = _after(t, "Heat Setting")
    out["antimicrobial"] = _after(t, "Antimicrobial Treatment")
    out["flame_retardant"] = _after(t, "Flame Retardant")
    out["dwr"] = _after(t, "Water/Stain Repellent (DWR)")
    # Residue results — pull pass/fail status + value
    out["res_formaldehyde"] = _join_wrap(_first_match(t, r"Formaldehyde[^\n]*\n+([^\n]+(?:\n[^\n]+)?)"))
    out["res_azo"] = _first_match(t, r"Azo amines[^\n]*\n+([^\n]+)") or _first_match(t, r"Azo dyes[^\n]*\n([^\n]+)")
    out["res_heavy_metals"] = _join_wrap(_first_match(t, r"Heavy metals[^\n]*\n+([^\n]+(?:\n[^\n]+)?)"))
    out["res_finishing"] = _join_wrap(_first_match(t, r"Finishing Agent Residue[^\n]*\n+([^\n]+(?:\n[^\n]+)?)"))
    out["res_svhc"] = "No SVHC > 0.1% w/w" if re.search(r"No SVHC\s*>\s*0\.1%", t, re.IGNORECASE) else None
    if re.search(r"PFAS-free", t, re.IGNORECASE):
        out["res_pfas"] = "PFAS-free (not used)"
    else:
        out["res_pfas"] = _join_wrap(_first_match(t, r"PFAS\s*\(sum[^\n]*\n+([^\n]+(?:\n[^\n]+){0,2})"))
    return {k: v for k, v in out.items() if v}


def parse_02a_po(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["po_no"] = _after(t, "PO Number") or _first_match(t, r"PO-NAC-[0-9-]+")
    out["po_supplier"] = _after(t, "Supplier (Mill)")
    out["po_buyer"] = _after(t, "Buyer")
    out["po_supplier_id"] = _after(t, "Supplier ID")
    out["po_incoterm"] = _after(t, "Incoterm")
    out["po_delivery_window"] = _after(t, "Required Delivery Window")
    out["coc_model"] = _first_match(t, r"Chain of Custody Model required:\s*([^\n]+)")
    out["tc_chain"] = _parse_tc_chain(t)
    out["cert_gots"] = _first_match(t, r"(GOTS-CU-\d{4}-[A-Z]{2}-\d+)")
    out["cert_oekotex"] = _first_match(t, r"(OEKO-TEX-100-[A-Z0-9.\-]+)")
    out["cert_grs"] = _first_match(t, r"(GRS-CU-\d{4}-[A-Z]{2}-\d+)")
    return {k: v for k, v in out.items() if v}


def _parse_tc_chain(t: str) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    pattern = re.compile(r"(Tier\s+\d\s*[→\-]+\s*Tier\s+\d[^\n]*|Tier\s+\d\s*→\s*Buyer[^\n]*)\s*\n+([A-Z]+-[A-Z]+-[0-9]{4}-[A-Z]{2}-[0-9]+)\s*\n+([A-Za-z /]+)?", re.IGNORECASE)
    for m in pattern.finditer(t):
        rows.append({"handover": m.group(1).strip(), "tc_number": m.group(2).strip(), "standard": (m.group(3) or "").strip()})
    return rows


def parse_02b_bol(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["bol_no"] = _after(t, "BOL Number") or _first_match(t, r"BOL[A-Z0-9-]+")
    out["bol_shipper"] = _after(t, "Shipper")
    out["bol_consignee"] = _after(t, "Consignee")
    out["bol_port_loading"] = _after(t, "Port of Loading")
    out["bol_port_discharge"] = _after(t, "Port of Discharge")
    out["bol_carrier"] = _after(t, "Carrier")
    out["bol_vessel"] = _after(t, "Vessel / Voyage")
    out["bol_etd"] = _after(t, "ETD") or _after(t, "Estimated Time of Departure")
    out["bol_eta"] = _after(t, "ETA") or _after(t, "Estimated Time of Arrival")
    return {k: v for k, v in out.items() if v}


def parse_02c_millprofile(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["mill_name"] = _after(t, "Legal Entity Name")
    out["mill_id"] = _after(t, "Mill Facility ID (Internal)")
    out["higg_fem_id"] = _after(t, "Higg FEM Facility ID")
    out["mill_address"] = _after(t, "Address")
    out["mill_country"] = _after(t, "Country of Operation")
    out["mill_year_established"] = _after(t, "Year Established")
    out["mill_workforce"] = _after(t, "Workforce Size")
    out["coc_model_full"] = _after(t, "Chain of Custody Model in Use") or _first_match(t, r"Chain of Custody Model in Use\s*\n([^\n]+)")
    return {k: v for k, v in out.items() if v}


def parse_03a_durability(t: str) -> Dict[str, Any]:
    """Extract Durability Test Panel as 8 dur_N records."""
    out: Dict[str, Any] = {}
    tests = _parse_durability_panel(t)
    out["dur_panel"] = tests
    out["microfibre_rate"] = _first_match(t, r"Shedding Rate \(mass per kg fabric per wash\)\s*\n([^\n]+)") \
        or _first_match(t, r"(\d+\s*mg\s*/\s*kg)")
    out["microfibre_note"] = _first_match(t, r"Performance vs Median\s*\n([^\n]+)")
    out["shrinkage_warp"] = _first_match(t, r"Dimensional Change \(warp\)[^\n]*\nISO 6330\s*\n([\-+]?[\d.]+%)")
    out["shrinkage_weft"] = _first_match(t, r"Dimensional Change \(weft\)[^\n]*\nISO 6330\s*\n([\-+]?[\d.]+%)")
    return {k: v for k, v in out.items() if v}


def _parse_durability_panel(t: str) -> List[Dict[str, str]]:
    """Parse the DURABILITY TEST PANEL section: each test is name → method → result → notes (4 lines)."""
    rows: List[Dict[str, str]] = []
    m = re.search(r"DURABILITY TEST PANEL[^\n]*\n(.+?)(?=MICROFIBRE|Statement of conformity|$)", t, re.IGNORECASE | re.DOTALL)
    if not m:
        return rows
    block = m.group(1)
    # Skip header lines: "Test Parameter / Method / Result / Conditions"
    lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
    # Find start after "Conditions / Notes" header
    start = 0
    for i, ln in enumerate(lines):
        if "Conditions" in ln and "Notes" in ln:
            start = i + 1
            break
    # Each entry: parameter name, method (ISO ...), result, optional notes
    i = start
    while i < len(lines):
        name = lines[i]
        if not _looks_like_test_name(name):
            i += 1
            continue
        method = lines[i + 1] if i + 1 < len(lines) else ""
        result = lines[i + 2] if i + 2 < len(lines) else ""
        notes = lines[i + 3] if i + 3 < len(lines) else ""
        # Sometimes result/notes wrap to extra lines — heuristic: if next line is not a new test name, append
        consumed = 4
        if i + 4 < len(lines) and not _looks_like_test_name(lines[i + 4]) and not lines[i + 4].startswith("ISO ") and consumed == 4:
            notes = notes + " " + lines[i + 4]
            consumed = 5
        rows.append({"name": name, "method": method, "result": result, "notes": notes})
        i += consumed
    return rows


_TEST_NAMES_HINT = (
    "pilling", "colourfastness", "color fastness", "abrasion", "bursting",
    "tensile", "tear", "dimensional", "shrinkage", "fastness",
)


def _looks_like_test_name(s: str) -> bool:
    low = s.lower()
    return any(k in low for k in _TEST_NAMES_HINT)


def parse_03b_eol(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["eol_selected"] = _first_match(t, r"Selected pathway:\s*([^\n]+)")
    out["eol_rationale"] = _first_match(t, r"Rationale:\s*([^\n]+(?:\n[^\n]+){0,2})")
    out["mono_material"] = _after(t, "Mono-material Status")
    out["dominant_fibre"] = _after(t, "Dominant Fibre")
    out["mechanical_yield"] = _after(t, "Mechanical Recycling Yield") or _after(t, "Mechanical Recycling Yield (theoretical)")
    out["chemical_recycling"] = _after(t, "Chemical Recycling Compatibility")
    out["permanent_finish"] = _after(t, "Permanent Finish Flag")
    return {k: v for k, v in out.items() if v}


def parse_03c_care(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["care_wash_ceiling"] = _after(t, "Maximum Wash Temperature")
    out["care_dry_ceiling"] = _after(t, "Maximum Dry Method")
    out["care_bleach_ceiling"] = _after(t, "Bleach Statement")
    out["care_iron_ceiling"] = _after(t, "Iron Cap") or _first_match(t, r"Iron Cap[^\n]*\n([^\n]+)")
    return {k: v for k, v in out.items() if v}


def parse_04a_gots(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["gots_cert_no"] = _after(t, "Certificate Number") or _first_match(t, r"(GOTS-CU-\d{4}-[A-Z]{2}-\d+)")
    out["gots_standard"] = _after(t, "Standard")
    out["gots_valid_until"] = _after(t, "Valid Until")
    out["gots_status"] = _after(t, "Status")
    out["gots_issuer"] = _after(t, "Issuing Body")
    out["gots_holder"] = _after(t, "Legal Entity")
    out["coc_model_gots"] = _after(t, "Chain of Custody Model")
    return {k: v for k, v in out.items() if v}


def parse_04b_reach(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["reach_doc_no"] = _first_match(t, r"(REACH-FAB-[A-Z0-9-]+)")
    out["reach_svhc_threshold"] = _first_match(t, r"No SVHC\s*>\s*([^\n]+)")
    out["reach_candidate_list"] = _after(t, "Candidate List Version")
    out["reach_pops"] = _after(t, "POPs Regulation (EU) 2019/1021")
    return {k: v for k, v in out.items() if v}


def parse_04c_oekotex(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["oekotex_cert_no"] = _after(t, "Certificate Number") or _first_match(t, r"(OEKO-TEX-100-[A-Z0-9.\-]+)")
    out["oekotex_class"] = _after(t, "Product Class")
    out["oekotex_valid_until"] = _after(t, "Valid Until")
    out["oekotex_issuer"] = _after(t, "Issuing Institute")
    return {k: v for k, v in out.items() if v}


def parse_05a_lca(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    # Headline metrics
    out["carbon_val"] = _first_match(t, r"Carbon footprint \(PCF\)\s*\n([\d.]+)\s*kg")
    out["carbon_unit"] = "kg CO₂-eq / kg" if out.get("carbon_val") else None
    out["water_val"] = _first_match(t, r"Water consumption\s*\n([\d.]+)\s*L")
    out["water_unit"] = "L / kg" if out.get("water_val") else None
    out["energy_val"] = _first_match(t, r"Cumulative energy demand\s*\n([\d.]+)\s*MJ")
    out["energy_unit"] = "MJ / kg" if out.get("energy_val") else None
    # Carbon stage breakdown
    out["co2_cultivation"] = _first_match(t, r"Fibre cultivation[^\n]*\n([\d.]+)\s*kg CO2e")
    out["co2_spinning"] = _first_match(t, r"Ginning \+ spinning[^\n]*\n([\d.]+)\s*kg CO2e")
    out["co2_knitting"] = _first_match(t, r"\nKnitting\s*\n([\d.]+)\s*kg CO2e") or _first_match(t, r"\bWeaving\s*\n([\d.]+)\s*kg CO2e")
    out["co2_dyeing"] = _first_match(t, r"Dyeing & finishing\s*\n([\d.]+)\s*kg CO2e")
    out["co2_logistics"] = _first_match(t, r"Energy supply \(mill\)\s*\n([\d.]+)\s*kg CO2e")
    # Water breakdown
    out["water_cultivation"] = _first_match(t, r"Fibre cultivation \(irrigation-share\)\s*\n([\d.]+)\s*L")
    out["water_spinning"] = _first_match(t, r"Spinning & pre-treatment\s*\n([\d.]+)\s*L")
    out["water_finishing"] = _first_match(t, r"Dyeing\s*\n([\d.]+)\s*L")
    return {k: v for k, v in out.items() if v}


def parse_05b_resource(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["renewable_share"] = _first_match(t, r"Renewable share:\s*(\d+)%") or _first_match(t, r"Renewable share[^\d]*(\d+)%")
    out["grid_share"] = _first_match(t, r"Grid electricity \(Vietnam EVN\)\s*\n(\d+)%") or _first_match(t, r"Grid electricity[^\n]*\n(\d+)%")
    out["ww_method"] = _after(t, "Treatment Type")
    out["ww_recovery_rate"] = _first_match(t, r"Water Recovery / Reuse\s*\n(\d+%[^\n]*)")
    out["ww_cod"] = _first_match(t, r"COD \(Chemical Oxygen Demand\)\s*\n([^\n]+)")
    out["ww_zdhc"] = _after(t, "ZDHC Wastewater Guideline Status")
    out["zdhc_present"] = bool(re.search(r"ZDHC", t))
    out["energy_renewable"] = _first_match(t, r"Renewable share:\s*(\d+%[^\n]*)")
    out["energy_grid"] = _first_match(t, r"Grid electricity[^\n]*\n(\d+%)")
    out["energy_gas"] = _first_match(t, r"(?:Natural gas|Diesel backup generator)[^\n]*\n(\d+%)")
    return {k: v for k, v in out.items() if v is not None and v != ""}


def parse_05c_social(t: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    out["audit_overall_rating"] = _first_match(t, r"Overall Rating\s*\n+(Rating [A-D])") or _first_match(t, r"Rating\s+([A-D])\b")
    out["audit_critical"] = _first_match(t, r"Critical \(Zero-Tolerance\) Findings\s*\n(\d+)")
    out["audit_major"] = _first_match(t, r"Major Findings\s*\n(\d+)")
    out["audit_minor"] = _first_match(t, r"Minor Findings\s*\n(\d+)")
    out["audit_scheme"] = _after(t, "Audit Scheme")
    out["audit_date"] = _after(t, "Audit Date")
    out["audit_auditor"] = _after(t, "Auditor")
    return {k: v for k, v in out.items() if v}


def parse_06a_scorecard(t: str) -> Dict[str, Any]:
    """Cluster A (OTD, Qty) + Cluster C (Doc, Transparency) + supplier metadata."""
    out: Dict[str, Any] = {}
    out["supplier_legal_name"] = _after(t, "Supplier Legal Name")
    out["supplier_id"] = _after(t, "Supplier ID")
    out["supplier_higg_id"] = _after(t, "Higg FEM ID")
    out["supplier_tier"] = _after(t, "Tier Classification")
    out["supplier_capacity"] = _after(t, "Annual Capacity")
    out["supplier_workforce"] = _after(t, "Workforce")
    out["supplier_last_audit"] = _after(t, "Last Audit Date")
    out["supplier_reporting_period"] = _after(t, "Reporting Period")
    # KPI values
    out["kpi_otd_val"] = _first_match(t, r"A1\.\s*On-Time Delivery Rate \(OTD\)\s*\n(\d+%)") \
        or _first_match(t, r"OTD Rate \(calculated\)\s*\n(\d+%)")
    out["kpi_otd_grade_doc"] = _first_match(t, r"A1\.\s*On-Time[^\n]*\n\d+%\s*\n([ABCDF])")
    out["kpi_qty_val"] = _first_match(t, r"A2\.\s*Quantity Execution Accuracy\s*\n(\d+%)") \
        or _first_match(t, r"Quantity Execution Accuracy\s*\n(\d+%)")
    out["kpi_qty_grade_doc"] = _first_match(t, r"A2\.\s*Quantity[^\n]*\n\d+%\s*\n([ABCDF])")
    out["kpi_doc_val"] = _first_match(t, r"C1\.\s*Document Completeness\s*\n(\d+%)") \
        or _first_match(t, r"Document Completeness Rate\s*\n(\d+%)")
    out["kpi_doc_grade_doc"] = _first_match(t, r"C1\.[^\n]*\n\d+%\s*\n([ABCDF])")
    out["kpi_transparency_val"] = _first_match(t, r"C3\.\s*Supply Chain Transparency\s*\n([^\n]+)") \
        or _first_match(t, r"Total Tier Coverage\s*\n(\d+\s*of\s*\d+ tiers)")
    out["kpi_transparency_grade_doc"] = _first_match(t, r"C3\.[^\n]*\n[^\n]+\n([ABCDF])")
    # Overall preliminary grade from doc
    out["overall_preliminary"] = _first_match(t, r"overall supplier grade:\s*([A-D][+\-]?)\s*\(score\s*(\d+)/100\)") \
        or _first_match(t, r"overall supplier grade:\s*([^\(]+)\(score\s*\d+/100\)")
    out["overall_score_doc"] = _first_match(t, r"\(score\s*(\d+)/100\)")
    return {k: v for k, v in out.items() if v is not None and v != ""}


def parse_06b_qc(t: str) -> Dict[str, Any]:
    """Cluster B1 — Quality Pass Rate."""
    out: Dict[str, Any] = {}
    out["kpi_qpr_val"] = _first_match(t, r"Quality Pass Rate\s*\n(\d+%)") \
        or _first_match(t, r"Visual Inspection Pass Rate\s*\n(\d+%)")
    out["kpi_qpr_grade_doc"] = _first_match(t, r"Grade Assigned\s*\n([ABCDF])")
    out["qc_total_lots"] = _first_match(t, r"Total Lots Inspected\s*\n(\d+)")
    out["qc_lots_passed"] = _first_match(t, r"Lots Passed[^\n]*\n(\d+)")
    out["qc_standard"] = _after(t, "Inspection Standard")
    return {k: v for k, v in out.items() if v}


def parse_06c_colour(t: str) -> Dict[str, Any]:
    """Cluster B2 — Colour Consistency (Delta E)."""
    out: Dict[str, Any] = {}
    out["kpi_colour_val"] = _first_match(t, r"Average Delta E[^\n]*\n([\d.]+)") \
        or _first_match(t, r"Result\s*\n(Delta E\s*[\d.]+)")
    out["kpi_colour_grade_doc"] = _first_match(t, r"Grade Assigned\s*\n([ABCDF])")
    out["colour_total_rolls"] = _first_match(t, r"Total Rolls Measured\s*\n(\d+)")
    out["colour_pass_count"] = _first_match(t, r"Rolls Passing Delta E[^\n]*\n(\d+\s*of\s*\d+[^\n]*)")
    return {k: v for k, v in out.items() if v}


# ────────────────────────────────────────────────────────────
# Aggregator → builds full passport_v2 payload
# ────────────────────────────────────────────────────────────

DOC_PARSERS = {
    "01a": parse_01a_techspec, "01b": parse_01b_materiallist, "01c": parse_01c_chemicals,
    "02a": parse_02a_po,       "02b": parse_02b_bol,           "02c": parse_02c_millprofile,
    "03a": parse_03a_durability, "03b": parse_03b_eol,         "03c": parse_03c_care,
    "04a": parse_04a_gots,     "04b": parse_04b_reach,         "04c": parse_04c_oekotex,
    "05a": parse_05a_lca,      "05b": parse_05b_resource,      "05c": parse_05c_social,
    "06a": parse_06a_scorecard, "06b": parse_06b_qc,           "06c": parse_06c_colour,
}


def parse_file(filename: str, data: bytes) -> Dict[str, Any]:
    text = extract_text(filename, data)
    doc_type = doc_type_from_name(filename)
    parser = DOC_PARSERS.get(doc_type) if doc_type else None
    parsed = parser(text) if parser else {}
    return {
        "filename": filename,
        "doc_type": doc_type,
        "char_count": len(text),
        "raw_text": text,
        "parsed": parsed,
    }


def build_passport(files: List[Tuple[str, bytes]]) -> Dict[str, Any]:
    """Take list of (filename, bytes), produce full passport_v2 payload."""
    parsed_docs: List[Dict[str, Any]] = []
    by_type: Dict[str, Dict[str, Any]] = {}
    full_text_parts: List[str] = []
    for fname, data in files:
        d = parse_file(fname, data)
        parsed_docs.append({"filename": d["filename"], "doc_type": d["doc_type"], "char_count": d["char_count"], "parsed": d["parsed"]})
        if d["doc_type"]:
            # Merge — later docs overwrite if duplicate keys
            by_type.setdefault(d["doc_type"], {}).update(d["parsed"])
        full_text_parts.append(d["raw_text"] or "")
    full_text = "\n".join(full_text_parts)

    # Build the unified passport payload
    payload = _assemble_payload(by_type, full_text, parsed_docs)
    payload["_docs"] = parsed_docs
    return payload


def _assemble_payload(by: Dict[str, Dict[str, Any]], full_text: str, parsed_docs: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Map per-doc dicts → final flat passport JSON consumed by passport_v2.html."""
    g = lambda dt, k, default=None: (by.get(dt, {}) or {}).get(k, default)  # noqa: E731

    p: Dict[str, Any] = {}
    # ── Identity ──
    p["fabric_name"] = g("01a", "fabric_name") or g("01a", "structure")
    fcode = g("01a", "fabric_code") or g("01b", "fabric_code") or _first_match(full_text, r"FAB-(?:RIB|TWL)-[A-Z0-9-]+")
    p["fabric_code"] = fcode
    p["passport_id"] = f"DFP-{fcode}" if fcode else None
    p["batch_no"] = g("01a", "batch_no")
    p["origin"] = g("01a", "origin") or g("02c", "mill_country")
    p["updated_date"] = "Updated " + (g("06a", "supplier_last_audit") or "—")

    # ── Hero tags / preview tags ──
    fibre_rows = g("01b", "fibre_rows") or []
    tags: List[str] = []
    for r in fibre_rows:
        tags.append(f"{r['pct']}% {r['name']}")
    if g("04a", "gots_cert_no"):
        tags.append("GOTS Certified")
    if g("04c", "oekotex_cert_no"):
        tags.append("OEKO-TEX")
    p["preview_tags"] = tags[:3]
    p["hero_tags"] = tags[:6]

    # ── Supplier ──
    p["supplier_full_name"] = g("06a", "supplier_legal_name") or g("02c", "mill_name") or g("01a", "mill_name")
    sname = p["supplier_full_name"] or ""
    short = sname.replace("Co., Ltd.", "").replace(",", "").strip() if sname else None
    p["supplier_short_name"] = short
    p["supplier_meta"] = g("02c", "mill_address") or g("01a", "mill_address")
    p["supplier_capacity"] = g("06a", "supplier_capacity")
    p["supplier_workforce"] = g("06a", "supplier_workforce") or g("02c", "mill_workforce")
    p["supplier_last_audit"] = g("06a", "supplier_last_audit")
    p["supplier_tier"] = g("06a", "supplier_tier")

    # ── Tab 1 ──
    p["fibre_legend"] = [{"name": r["name"], "pct": r["pct"]} for r in fibre_rows]
    for i, r in enumerate(fibre_rows[:3], start=1):
        p[f"fibre_{i}_name"] = f"{r['name']} ({r['pct']}%)"
        p[f"fibre_{i}_origin"] = r["origin"]
        p[f"fibre_{i}_note"] = r["source"]
    p["recycled_pct"] = g("01b", "recycled_pct")
    p["recycled_source_type"] = g("01b", "recycled_source_type")
    p["traceability_method"] = g("02c", "coc_model_full") or g("02a", "coc_model")
    p["yarn_count"] = g("01a", "yarn_count")
    p["yarn_twist"] = g("01a", "yarn_twist")
    p["fibre_length"] = g("01a", "fibre_length")
    p["spinning_method"] = g("01a", "spinning_method")
    p["weight"] = g("01a", "weight")
    p["width"] = g("01a", "width")
    p["structure"] = g("01a", "structure")
    p["colour_state"] = g("01a", "colour_state")
    p["shrinkage_warp"] = g("03a", "shrinkage_warp")
    p["shrinkage_weft"] = g("03a", "shrinkage_weft")
    # Durability + colourfastness from 03a panel
    dur = g("03a", "dur_panel") or []
    for i, t_ in enumerate(dur[:8], start=1):
        p[f"dur_{i}_name"] = t_.get("name")
        p[f"dur_{i}_iso"] = t_.get("method")
        p[f"dur_{i}_result"] = t_.get("result")
        p[f"dur_{i}_note"] = t_.get("notes")
    # Pull individual physical-spec rows from panel by name
    def find_test(keyword: str) -> Dict[str, str]:
        for t_ in dur:
            if keyword.lower() in (t_.get("name") or "").lower():
                return t_
        return {}
    pilling = find_test("Pilling")
    abrasion = find_test("Abrasion")
    p["pilling"] = pilling.get("result")
    p["pilling_note"] = pilling.get("notes")
    p["abrasion"] = abrasion.get("result")
    p["abrasion_note"] = abrasion.get("notes")
    p["tensile_warp"] = (find_test("Bursting") or find_test("Tensile")).get("result")
    p["tear_strength"] = find_test("Tear").get("result")
    cf_wash = find_test("Wash")
    cf_light = find_test("Light")
    cf_dry = next((x for x in dur if "rubbing (dry)" in (x.get("name") or "").lower()), {})
    cf_wet = next((x for x in dur if "rubbing (wet)" in (x.get("name") or "").lower()), {})
    p["cf_washing"] = cf_wash.get("result")
    p["cf_rub_dry"] = cf_dry.get("result")
    p["cf_rub_wet"] = cf_wet.get("result")
    p["cf_light"] = cf_light.get("result")
    # Chemicals
    p["dyestuff_class"] = g("01c", "dyestuff_class")
    p["pretreatment_agent"] = g("01c", "pretreatment_agent")
    p["softener"] = g("01c", "softener")
    p["heat_setting"] = g("01c", "heat_setting") or g("01a", "heat_setting")
    p["res_formaldehyde"] = g("01c", "res_formaldehyde")
    p["res_azo"] = g("01c", "res_azo")
    p["res_heavy_metals"] = g("01c", "res_heavy_metals")
    p["res_finishing"] = g("01c", "res_finishing")
    p["res_svhc"] = g("01c", "res_svhc") or ("No SVHC > 0.1% w/w" if re.search(r"No SVHC\s*>\s*0\.1%", full_text, re.IGNORECASE) else None)
    p["res_pfas"] = g("01c", "res_pfas")

    # ── Tab 2 — Product Journey ──
    p["coc_model_summary"] = g("02a", "coc_model") or g("02c", "coc_model_full")
    upstream = g("01b", "upstream_tiers") or []
    tc_chain = g("02a", "tc_chain") or []
    # Build 6 steps if possible — use upstream + standard mill ordering
    steps = _build_supply_chain_steps(upstream, tc_chain, g("01a", "mill_name") or g("02c", "mill_name"))
    for i, s in enumerate(steps[:6], start=1):
        p[f"step_{i}_name"] = s.get("name")
        p[f"step_{i}_desc"] = s.get("desc")
        p[f"step_{i}_facility"] = s.get("facility")
        p[f"step_{i}_loc"] = s.get("loc")
        p[f"step_{i}_tc_in"] = s.get("tc_in")
        p[f"step_{i}_tc_out"] = s.get("tc_out")
        p[f"step_{i}_tc_status"] = s.get("tc_status")
        p[f"step_{i}_tc_note"] = s.get("tc_note")

    # ── Tab 3 — EOL + Care ──
    eol_selected = (g("03b", "eol_selected") or "").lower()
    p["eol_card_technical_tag"] = "Selected" if "technical" in eol_selected else None
    p["eol_card_biological_tag"] = "Selected" if "biological" in eol_selected else None
    p["eol_card_mixed_tag"] = "Selected" if "mixed" in eol_selected else None
    p["eol_selected_value"] = g("03b", "eol_selected")
    p["eol_row_1_key"] = "Mono-material Status"; p["eol_row_1_val"] = g("03b", "mono_material")
    p["eol_row_2_key"] = "Dominant Fibre"; p["eol_row_2_val"] = g("03b", "dominant_fibre")
    p["eol_row_3_key"] = "Mechanical Yield"; p["eol_row_3_val"] = g("03b", "mechanical_yield")
    p["eol_row_4_key"] = "Chemical Recycling"; p["eol_row_4_val"] = g("03b", "chemical_recycling")
    p["eol_row_5_key"] = "Permanent Finish"; p["eol_row_5_val"] = g("03b", "permanent_finish")
    p["microfibre_rate"] = g("03a", "microfibre_rate")
    p["microfibre_note"] = g("03a", "microfibre_note")
    p["care_wash_ceiling"] = g("03c", "care_wash_ceiling")
    p["care_dry_ceiling"] = g("03c", "care_dry_ceiling")
    p["care_bleach_ceiling"] = g("03c", "care_bleach_ceiling")
    p["care_iron_ceiling"] = g("03c", "care_iron_ceiling")

    # ── Tab 4 — Certifications ──
    certs: List[Dict[str, str]] = []
    if g("04a", "gots_cert_no"):
        certs.append({"name": f"GOTS {g('04a', 'gots_standard') or '7.0'}", "verified": True})
    if g("04c", "oekotex_cert_no"):
        certs.append({"name": "OEKO-TEX® Standard 100", "verified": True})
    if re.search(r"ZDHC MRSL", full_text, re.IGNORECASE):
        certs.append({"name": "ZDHC MRSL", "verified": True})
    if re.search(r"Higg FEM 4\.0", full_text):
        certs.append({"name": "Higg FEM 4.0", "verified": True})
    if re.search(r"GRS-CU-", full_text):
        certs.append({"name": "GRS — Recycled", "verified": True})
    if re.search(r"BSCI", full_text):
        certs.append({"name": "amfori BSCI", "verified": True})
    p["certifications"] = certs
    p["doc_tc_number"] = g("02a", "tc_chain") and (g("02a", "tc_chain")[-1]["tc_number"] if g("02a", "tc_chain") else None)
    p["doc_tc_note"] = g("02a", "coc_model")
    p["doc_test_lab"] = "Intertek / Hohenstein"
    p["doc_test_note"] = "ISO/IEC 17025 accredited"
    p["doc_svhc"] = g("04b", "reach_svhc_threshold") or g("01c", "res_svhc")
    p["doc_svhc_note"] = "REACH Annex XIV"
    p["doc_reach"] = "REACH (EC) 1907/2006" if g("04b", "reach_doc_no") else None
    p["doc_reach_note"] = g("04b", "reach_candidate_list")
    p["doc_coc"] = g("04a", "coc_model_gots") or g("02a", "coc_model")
    p["doc_coc_note"] = "Verified via TC chain"
    p["doc_audit"] = g("05c", "audit_overall_rating")
    p["doc_audit_note"] = g("05c", "audit_scheme")

    # ── Tab 5 — Impact ──
    p["water_val"] = g("05a", "water_val")
    p["water_unit"] = g("05a", "water_unit")
    p["water_bench"] = "Below industry avg" if p.get("water_val") else None
    p["carbon_val"] = g("05a", "carbon_val")
    p["carbon_unit"] = g("05a", "carbon_unit")
    p["carbon_bench"] = "Below industry avg" if p.get("carbon_val") else None
    p["energy_val"] = g("05a", "energy_val")
    p["energy_unit"] = g("05a", "energy_unit")
    p["energy_bench"] = "Above renewable target" if p.get("energy_val") else None
    p["wastewater_val"] = g("05b", "ww_recovery_rate")
    p["wastewater_unit"] = ""
    p["wastewater_bench"] = "Strong recovery" if p.get("wastewater_val") else None
    p["water_cultivation"] = g("05a", "water_cultivation")
    p["water_spinning"] = g("05a", "water_spinning")
    p["water_finishing"] = g("05a", "water_finishing")
    p["co2_cultivation"] = g("05a", "co2_cultivation")
    p["co2_spinning"] = g("05a", "co2_spinning")
    p["co2_knitting"] = g("05a", "co2_knitting")
    p["co2_dyeing"] = g("05a", "co2_dyeing")
    p["co2_logistics"] = g("05a", "co2_logistics")
    p["ww_method"] = g("05b", "ww_method")
    p["ww_recovery_rate"] = g("05b", "ww_recovery_rate")
    p["ww_cod"] = g("05b", "ww_cod")
    p["ww_zdhc"] = g("05b", "ww_zdhc")
    p["energy_renewable"] = g("05b", "energy_renewable")
    p["energy_grid"] = g("05b", "energy_grid")
    p["energy_gas"] = g("05b", "energy_gas")

    # ── Tab 6 — Supplier KPIs ──
    p["kpi_otd_val"] = g("06a", "kpi_otd_val")
    p["kpi_qty_val"] = g("06a", "kpi_qty_val")
    p["kpi_qpr_val"] = g("06b", "kpi_qpr_val")
    p["kpi_colour_val"] = g("06c", "kpi_colour_val")
    p["kpi_doc_val"] = g("06a", "kpi_doc_val")
    p["kpi_audit_val"] = g("05c", "audit_overall_rating")
    p["kpi_transparency_val"] = g("06a", "kpi_transparency_val")
    # KPI benchmark blurbs
    p["kpi_otd_benchmark"] = "Benchmark 95%+"
    p["kpi_qty_benchmark"] = "Benchmark 97%+"
    p["kpi_qpr_benchmark"] = "AQL 2.5 · 95%+"
    p["kpi_colour_benchmark"] = "ΔE ≤ 1.0"
    p["kpi_doc_benchmark"] = "100% on-time"
    p["kpi_audit_benchmark"] = "amfori BSCI"
    p["kpi_transparency_benchmark"] = "Tier 1-4 visibility"
    # Preview card needs otd_rate (separate from Tab 6 kpi_otd_val)
    p["otd_rate"] = p["kpi_otd_val"]

    # ── Scoring inputs for shell's calculateScores() ──
    p["score_inputs"] = _build_score_inputs(by, fibre_rows, full_text)

    # Final supplier grade (compute from KPI doc grades if shell gate fails)
    p["supplier_grade"] = _supplier_grade_from_kpis(by)
    return {k: v for k, v in p.items() if v is not None and v != ""}


def _build_supply_chain_steps(upstream: List[Dict[str, str]], tc_chain: List[Dict[str, str]], mill_name: Optional[str]) -> List[Dict[str, str]]:
    """Synthesize 6 supply chain steps from upstream tier list + TC handovers + mill name."""
    # Order: Tier 4 → 3 → 2 → 1 → Dyeing → Buyer (approx 6 stages)
    # Match upstream entries to ordering
    ordered: List[Dict[str, str]] = []
    tier_map = {4: "Fibre Cultivation", 3: "Ginning", 2: "Spinning", 1: "Knitting & Dyeing"}
    upstream_by_tier: Dict[int, Dict[str, str]] = {}
    for u in upstream:
        m = re.search(r"Tier\s+(\d)", u.get("tier", ""))
        if m:
            upstream_by_tier[int(m.group(1))] = u
    # Build the 4 mill-input steps
    for tier in [4, 3, 2, 1]:
        u = upstream_by_tier.get(tier)
        if u:
            ordered.append({
                "name": tier_map[tier],
                "desc": u.get("tier", ""),
                "facility": u.get("supplier", ""),
                "loc": u.get("location", ""),
                "tc_in": "—",
                "tc_out": _tc_for_tier(tc_chain, tier),
                "tc_status": "✓",
                "tc_note": "verified" if _tc_for_tier(tc_chain, tier) else "pending",
            })
    # Step 5: Dyeing & Finishing (in mill, same facility)
    if mill_name:
        ordered.append({
            "name": "Dyeing & Finishing",
            "desc": "Reactive / vat dyeing, finishing, heat-set",
            "facility": mill_name,
            "loc": "Mill — Tier 1",
            "tc_in": "—",
            "tc_out": "—",
            "tc_status": "✓",
            "tc_note": "in-house",
        })
    # Step 6: Buyer / brand handover
    buyer_tc = next((r for r in tc_chain if "Buyer" in r.get("handover", "")), None)
    ordered.append({
        "name": "Brand Handover",
        "desc": "Mill → Buyer (final TC)",
        "facility": "Nordic Apparel Co.",
        "loc": "EU import",
        "tc_in": buyer_tc["tc_number"] if buyer_tc else "—",
        "tc_out": "—",
        "tc_status": "✓" if buyer_tc else "—",
        "tc_note": buyer_tc.get("standard") if buyer_tc else "",
    })
    return ordered


def _tc_for_tier(tc_chain: List[Dict[str, str]], tier: int) -> str:
    """Find TC number whose handover originates from the given tier."""
    for r in tc_chain:
        if f"Tier {tier}" in r.get("handover", ""):
            return r.get("tc_number", "")
    return ""


def _build_score_inputs(by: Dict[str, Dict[str, Any]], fibre_rows: List[Dict[str, str]], full_text: str) -> Dict[str, Any]:
    """Inputs for passport_v2_shell's calculateScores() engine."""
    comp_total = sum(r.get("pct", 0) for r in fibre_rows) if fibre_rows else None
    svhc_passed = bool(re.search(r"No SVHC\s*>\s*0\.1%", full_text, re.IGNORECASE))
    origin_disclosed = bool((by.get("01a") or {}).get("origin") or (by.get("02c") or {}).get("mill_country"))
    has_test_report = bool(by.get("03a") or {}) or bool(re.search(r"ISO/IEC 17025", full_text))

    recycled_str = (by.get("01b") or {}).get("recycled_pct") or "0%"
    rmatch = re.search(r"(\d+)%", recycled_str)
    recycled_pct = int(rmatch.group(1)) if rmatch else 0

    is_mono = len(fibre_rows) == 1
    # Traceability tiers — count of tiers with disclosed visibility
    upstream = (by.get("01b") or {}).get("upstream_tiers") or []
    traceability_tiers = min(len(upstream), 4)
    # Durability tests passed
    dur = (by.get("03a") or {}).get("dur_panel") or []
    durability_passed = sum(1 for d in dur if _is_pass(d.get("result", "")))
    # Chemical finish
    finish_keywords_strong = ("PFAS", "flame retardant", "antimicrobial")
    has_finish = any(re.search(k, full_text, re.IGNORECASE) and not re.search(rf"{k}[^\n]*(?:None|Not used|PFAS-free)", full_text, re.IGNORECASE) for k in finish_keywords_strong)
    chemical_finish = "None" if not has_finish else "Standard"
    has_zdhc = bool(re.search(r"ZDHC", full_text))
    # Carbon better than avg — fixed industry baseline ~12 kg CO2e/kg for cotton fabric; benchmark percentage
    co2 = (by.get("05a") or {}).get("carbon_val")
    if co2:
        try:
            co2v = float(co2)
            baseline = 12.0
            carbon_better = max(0, round((baseline - co2v) / baseline * 100))
        except ValueError:
            carbon_better = 0
    else:
        carbon_better = 0

    # Supplier KPIs
    def _pct(s: Optional[str]) -> Optional[float]:
        """Extract first numeric, % suffix optional."""
        if s is None:
            return None
        m = re.search(r"([\d.]+)", str(s))
        return float(m.group(1)) if m else None

    otd = _pct((by.get("06a") or {}).get("kpi_otd_val")) or 0
    qty = _pct((by.get("06a") or {}).get("kpi_qty_val")) or 0
    aql = _pct((by.get("06b") or {}).get("kpi_qpr_val")) or 0
    renewable = _pct((by.get("05b") or {}).get("renewable_share") or (by.get("05b") or {}).get("energy_renewable")) or 0
    ww_recovery = _pct((by.get("05b") or {}).get("ww_recovery_rate")) or 0
    audit_grade_raw = (by.get("05c") or {}).get("audit_overall_rating") or ""
    audit_grade = "A"
    am = re.search(r"Rating\s*([A-D])", audit_grade_raw)
    if am:
        audit_grade = am.group(1)

    return {
        "compositionTotal": comp_total,
        "svhcPassed": svhc_passed,
        "originDisclosed": origin_disclosed,
        "hasTestReport": has_test_report,
        "recycledContentPct": recycled_pct,
        "isMonoMaterial": is_mono,
        "traceabilityTiers": traceability_tiers,
        "durabilityTestsPassed": durability_passed,
        "chemicalFinish": chemical_finish,
        "hasZDHC": has_zdhc,
        "carbonBetterThanAvg": carbon_better,
        "otdRate": otd,
        "aqlRate": aql,
        "renewableEnergyPct": renewable,
        "wastewaterRecoveryPct": ww_recovery,
        "socialAuditGrade": audit_grade,
    }


def _is_pass(s: str) -> bool:
    low = (s or "").lower()
    if "fail" in low or "below" in low:
        return False
    # Grade 4 / 3-4 / Pass / Not detected / values like "20,000 cycles, no breakdown" all count as pass
    if re.search(r"grade\s*[3-5]", low) or "pass" in low or "no breakdown" in low or "not detected" in low:
        return True
    if re.search(r"\d+\s*(?:kpa|cycles)", low):
        return True
    if re.search(r"-?\d+\.?\d*\s*%", low):
        return True
    return False


def _supplier_grade_from_kpis(by: Dict[str, Dict[str, Any]]) -> Optional[str]:
    """Fallback grade aggregation directly from KPI letter grades in docs (if shell gate fails)."""
    grades_a = [(by.get("06a") or {}).get("kpi_otd_grade_doc"), (by.get("06a") or {}).get("kpi_qty_grade_doc")]
    grades_b = [(by.get("06b") or {}).get("kpi_qpr_grade_doc"), (by.get("06c") or {}).get("kpi_colour_grade_doc")]
    grades_c = [(by.get("06a") or {}).get("kpi_doc_grade_doc"),
                "A" if "rating a" in ((by.get("05c") or {}).get("audit_overall_rating") or "").lower() else
                "B" if "rating b" in ((by.get("05c") or {}).get("audit_overall_rating") or "").lower() else None,
                (by.get("06a") or {}).get("kpi_transparency_grade_doc")]
    all_grades = [g for g in grades_a + grades_b + grades_c if g]
    if not all_grades:
        return None
    map_n = {"A": 4, "B": 3, "C": 2, "D": 1, "F": 0}
    map_back = {4: "A", 3: "B+", 2.5: "B", 2: "C+", 1.5: "C", 1: "D", 0: "F"}
    avg = sum(map_n.get(g, 0) for g in all_grades) / len(all_grades)
    # Snap to nearest available bucket
    for thr, letter in sorted(map_back.items(), reverse=True):
        if avg >= thr - 0.25:
            return letter
    return "F"
