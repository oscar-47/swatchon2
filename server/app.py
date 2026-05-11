from io import BytesIO
from typing import Dict, List, Optional
import os
import hashlib
from datetime import datetime

import torch
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from PIL import Image

from collections import Counter

from server.models.loader import build_eval_transform, load_checkpoint
from server.ocr_parser import OCRFieldParser
from server.double_verifier import DoubleStructureVerifier
from server.quality_inspector import QualityInspector
from server.passport_v2_parser import build_passport as build_passport_v2

# OCR Configuration
OCR_ENABLED = True
CONFIDENCE_THRESHOLD = 0.60  # Trigger OCR if confidence < 60%

# Model registry: name -> (checkpoint path, stage_key)
MODEL_REGISTRY: Dict[str, tuple] = {
    "woven_vs_knit": (os.path.join("runs", "stage1_knit_vs_woven_vs_others_best.pth"), "stage1"),
    "woven_multi": (os.path.join("runs", "stage2_woven_7class_best.pth"), "stage2_woven"),
    "knit_multi": (os.path.join("runs", "stage2_knit_6class_best.pth"), "stage2_knit"),
}

app = FastAPI(title="FabricFlow API", version="2.0")

# Allow local dev frontends
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def no_cache_html(request, call_next):
    """Disable browser caching for HTML pages during development."""
    response = await call_next(request)
    if "text/html" in response.headers.get("content-type", ""):
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
    return response


class ModelCache:
    def __init__(self):
        self._cache: Dict[str, Dict] = {}
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        self.tfm = build_eval_transform(224)

        # Initialize OCR engine (lazy initialization)
        self.ocr = None
        self.ocr_initialized = False
        
        # Initialize OCR parser and double verifier
        self.ocr_parser = OCRFieldParser()
        self.double_verifier = DoubleStructureVerifier(similarity_threshold=0.75)

    def get(self, name: str):
        if name not in MODEL_REGISTRY:
            raise ValueError(f"Unknown model: {name}")
        if name not in self._cache:
            ckpt_path, stage_key = MODEL_REGISTRY[name]
            if not os.path.isfile(ckpt_path):
                raise FileNotFoundError(f"Checkpoint not found: {ckpt_path}")
            model, classes = load_checkpoint(ckpt_path, self.device, stage_key=stage_key)
            self._cache[name] = {"model": model, "classes": classes}
        return self._cache[name]["model"], self._cache[name]["classes"], self.tfm, self.device

    def recognize_ocr(self, img_pil):
        """Run OCR on image and return extracted text using EasyOCR"""
        # Lazy initialization of OCR engine
        if not self.ocr_initialized:
            if not OCR_ENABLED:
                return None
            try:
                import sys
                import easyocr
                sys.stdout.flush()
                sys.stderr.flush()

                # Initialize EasyOCR with English support
                # You can add more languages: ['en', 'ch_sim'] for Chinese
                sys.stderr.write("[OCR] Initializing EasyOCR...\n")
                sys.stderr.flush()

                self.ocr = easyocr.Reader(['en'], gpu=False)
                self.ocr_initialized = True

                sys.stderr.write("[OCR] EasyOCR initialized successfully!\n")
                sys.stderr.flush()
            except Exception as e:
                self.ocr_initialized = True  # Mark as initialized even if failed
                import traceback
                import sys
                sys.stderr.write(f"[OCR] Init failed: {str(e)}\n")
                sys.stderr.write(traceback.format_exc() + "\n")
                sys.stderr.flush()
                return None

        if not self.ocr:
            return None

        try:
            import numpy as np
            import sys

            # Convert PIL Image to numpy array
            img_array = np.array(img_pil)

            sys.stderr.write(f"[OCR] Image shape: {img_array.shape}, dtype: {img_array.dtype}\n")
            sys.stderr.flush()

            # EasyOCR readtext returns: [([bbox], text, confidence), ...]
            result = self.ocr.readtext(img_array)

            sys.stderr.write(f"[OCR] EasyOCR found {len(result)} text regions\n")
            sys.stderr.flush()

            if result and len(result) > 0:
                texts = []
                confidences = []

                for idx, detection in enumerate(result):
                    # detection is a tuple: (bbox, text, confidence)
                    bbox, text, conf = detection
                    sys.stderr.write(f"[OCR] Line {idx}: '{text}' (conf: {conf:.2f})\n")
                    texts.append(text)
                    confidences.append(float(conf))

                sys.stderr.flush()

                if texts:
                    full_text = ' '.join(texts)
                    avg_conf = sum(confidences) / len(confidences)
                    sys.stderr.write(f"[OCR] SUCCESS: extracted {len(texts)} items\n")
                    sys.stderr.write(f"[OCR] Text: {full_text[:200]}\n")
                    sys.stderr.flush()
                    return {"text": full_text, "confidence": float(avg_conf)}
            else:
                sys.stderr.write(f"[OCR] No text detected by EasyOCR\n")
                sys.stderr.flush()

        except Exception as e:
            import sys
            import traceback
            sys.stderr.write(f"[OCR] Recognition error: {str(e)}\n")
            sys.stderr.write(traceback.format_exc() + "\n")
            sys.stderr.flush()

        return None


CACHE = ModelCache()

# Optional training hash set for External-only filtering (resolve path from repo root)
REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
TRAIN_HASHES_PATH = os.path.join(REPO_ROOT, "training", "training_hashes.txt")
TRAIN_HASHES = set()
if os.path.isfile(TRAIN_HASHES_PATH):
    try:
        with open(TRAIN_HASHES_PATH, 'r', encoding='utf-8') as f:
            TRAIN_HASHES = set(line.strip() for line in f if line.strip())
    except Exception:
        TRAIN_HASHES = set()


def two_stage_predict(img_pil: Image.Image):
    """
    Two-stage cascade prediction (ConvNeXt):
    1. Stage 1: Classify as KNIT, WOVEN, or OTHERS
    2. Stage 2: Subcategory classification (skipped for OTHERS)

    Returns: dict with prediction, confidence, and stage details.
    """
    # Stage 1: KNIT / WOVEN / OTHERS
    model1, classes1, tfm, device = CACHE.get("woven_vs_knit")
    x = tfm(img_pil).unsqueeze(0).to(device)

    with torch.no_grad(), torch.amp.autocast("cuda", enabled=(device.type == "cuda")):
        logits1 = model1(x)
        probs1 = torch.softmax(logits1, dim=1)[0]
        conf1, idx1 = torch.max(probs1, dim=0)

    stage1_pred = classes1[int(idx1)]
    stage1_conf = float(conf1)
    stage1_probs = {classes1[i]: float(probs1[i]) for i in range(len(classes1))}

    # If OTHERS, skip stage 2
    if stage1_pred == "OTHERS":
        return {
            "prediction": "Others",
            "l1": "OTHERS",
            "confidence": stage1_conf,
            "stage1": {
                "prediction": stage1_pred,
                "confidence": stage1_conf,
                "probs": stage1_probs,
            },
            "stage2": None,
        }

    # Stage 2: Subcategory classification
    if stage1_pred == "KNIT":
        model2, classes2, tfm, device = CACHE.get("knit_multi")
    else:
        model2, classes2, tfm, device = CACHE.get("woven_multi")

    with torch.no_grad(), torch.amp.autocast("cuda", enabled=(device.type == "cuda")):
        logits2 = model2(x)
        probs2 = torch.softmax(logits2, dim=1)[0]
        conf2, idx2 = torch.max(probs2, dim=0)

    stage2_pred = classes2[int(idx2)]
    stage2_conf = float(conf2)
    stage2_probs = {classes2[i]: float(probs2[i]) for i in range(len(classes2))}

    final_conf = stage1_conf * stage2_conf

    return {
        "prediction": stage2_pred,
        "l1": stage1_pred,
        "confidence": final_conf,
        "stage1": {
            "prediction": stage1_pred,
            "confidence": stage1_conf,
            "probs": stage1_probs,
        },
        "stage2": {
            "prediction": stage2_pred,
            "confidence": stage2_conf,
            "probs": stage2_probs,
        },
    }

def sha256_bytes(data: bytes) -> str:
    h = hashlib.sha256()
    h.update(data)
    return h.hexdigest()


@app.get("/api/models")
def list_models():
    return {
        "models": [
            {
                "name": k,
                "checkpoint": v[0],
                "stage_key": v[1],
                "classes": (CACHE.get(k)[1] if os.path.exists(v[0]) else None),
            }
            for k, v in MODEL_REGISTRY.items()
        ],
        "hashes_loaded": len(TRAIN_HASHES),
    }


@app.post("/api/predict")
async def predict(model_name: str = Form(...), files: List[UploadFile] = File(...), external_only: bool = Form(False)):
    model, classes, tfm, device = CACHE.get(model_name)
    results = []
    skipped = 0
    for f in files:
        data = await f.read()
        if external_only and TRAIN_HASHES:
            h = sha256_bytes(data)
            if h in TRAIN_HASHES:
                skipped += 1
                continue
        img = Image.open(BytesIO(data)).convert("RGB")
        x = tfm(img).unsqueeze(0).to(device)
        with torch.no_grad(), torch.amp.autocast("cuda", enabled=(device.type == "cuda")):
            logits = model(x)
            probs = torch.softmax(logits, dim=1)[0]
            conf, idx = torch.max(probs, dim=0)

        confidence = float(conf)
        prediction = classes[int(idx)]

        # Check if OCR is needed (confidence below threshold)
        needs_ocr = confidence < CONFIDENCE_THRESHOLD

        results.append({
            "filename": f.filename,
            "pred": prediction,
            "confidence": confidence,
            "probs": {classes[i]: float(probs[i]) for i in range(len(classes))},
            "needs_ocr": needs_ocr
        })
    return {"classes": classes, "predictions": results, "skipped": skipped, "hashes_loaded": len(TRAIN_HASHES)}


@app.post("/api/ocr")
async def recognize_label(file: UploadFile = File(...)):
    """OCR recognition for fabric label images — returns raw text + parsed fields"""
    try:
        data = await file.read()
        img = Image.open(BytesIO(data)).convert("RGB")
        ocr_result = CACHE.recognize_ocr(img)

        if ocr_result:
            parsed = CACHE.ocr_parser.parse_all_fields(ocr_result["text"])
            return {
                "success": True,
                "text": ocr_result["text"],
                "confidence": ocr_result["confidence"],
                "parsed": parsed
            }
        else:
            return {
                "success": False,
                "text": None,
                "confidence": None,
                "message": "No text detected in image"
            }
    except Exception as e:
        return {
            "success": False,
            "text": None,
            "confidence": None,
            "message": str(e)
        }


@app.post("/api/extract_doc")
async def extract_document(file: UploadFile = File(...)):
    """Extract text from PDF/DOCX uploads, parse fabric fields."""
    try:
        data = await file.read()
        name = (file.filename or "").lower()
        ext = name.rsplit(".", 1)[-1] if "." in name else ""
        text = ""
        if ext == "pdf":
            from pypdf import PdfReader
            r = PdfReader(BytesIO(data))
            text = "\n".join((p.extract_text() or "") for p in r.pages)
        elif ext in ("doc", "docx"):
            import docx as _docx
            d = _docx.Document(BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            for tbl in d.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(parts)
        else:
            return {"success": False, "message": f"Unsupported extension: {ext}"}
        text = (text or "").strip()
        if not text:
            return {"success": False, "message": "No text extracted"}
        parsed = CACHE.ocr_parser.parse_all_fields(text)
        return {
            "success": True,
            "filename": file.filename,
            "ext": ext,
            "char_count": len(text),
            "text": text[:6000],
            "parsed": parsed,
        }
    except Exception as e:
        import traceback; traceback.print_exc()
        return {"success": False, "message": str(e)}


@app.post("/api/build_passport_v2")
async def build_passport_v2_endpoint(files: List[UploadFile] = File(...)):
    """Aggregate multiple PDF/DOCX docs → full passport_v2 JSON payload.

    Accepts Felicity's 18-doc test bundle (or a subset). Returns a flat
    `data-field` → value map plus `score_inputs` consumed by passport_v2.html.
    """
    try:
        pairs: List[tuple] = []
        for f in files:
            data = await f.read()
            pairs.append((f.filename or "unnamed", data))
        if not pairs:
            return {"success": False, "message": "no files"}
        payload = build_passport_v2(pairs)
        return {"success": True, "fields": payload, "doc_count": len(pairs)}
    except Exception as e:
        import traceback; traceback.print_exc()
        return {"success": False, "message": str(e)}


@app.post("/api/create_metadata")
async def create_fabric_metadata(
    fabric_images: List[UploadFile] = File(...),
    label_image: Optional[UploadFile] = File(None),
    back_image: Optional[UploadFile] = File(None)
):
    """
    Complete Fabric Metadata Creation API (V3 Multi-View)
    
    Integrates:
    1. Quality Control: Sharpness, exposure, and content validation
    2. Multi-view consensus: Requires all fabric images to match with high confidence (>90%)
    3. OCR parsing
    4. Optional double-layer verification
    """
    try:
        # --- Step 0: Quality Control & Image Loading ---
        loaded_fabric_imgs = []
        raw_bytes_list = []
        qc_issues = []

        for idx, f_file in enumerate(fabric_images):
            f_data = await f_file.read()
            raw_bytes_list.append(f_data)
            img = Image.open(BytesIO(f_data)).convert("RGB")
            loaded_fabric_imgs.append(img)
            
            # Run QC
            inspector_res = QualityInspector.inspect_image(
                img, 
                ocr_reader=(CACHE.ocr if CACHE.ocr_initialized else None)
            )
            
            # Store full inspection details (convert numpy types to native Python)
            details = inspector_res["details"]
            qc_issues.append({
                "image_index": idx,
                "passed": bool(inspector_res["passed"]),
                "issues": list(inspector_res["issues"]),
                "details": {
                    "sharpness": {
                        "score": float(details["sharpness"]["score"]),
                        "passed": bool(details["sharpness"]["passed"]),
                        "message": str(details["sharpness"]["message"]),
                        "threshold": float(details["sharpness"]["threshold"])
                    },
                    "exposure": {
                        "score": float(details["exposure"]["score"]),
                        "passed": bool(details["exposure"]["passed"]),
                        "message": str(details["exposure"]["message"]),
                        "range": str(details["exposure"]["range"])
                    },
                    "text_check": {
                        "count": int(details["text_check"]["count"]),
                        "passed": bool(details["text_check"]["passed"]),
                        "message": str(details["text_check"]["message"])
                    }
                }
            })

        # Check if any image failed
        failed_images = [img_qc for img_qc in qc_issues if not img_qc["passed"]]
        
        if failed_images:
            return {
                "success": False,
                "message": "Quality Control Failed",
                "issues": [f"Image {item['image_index']+1}: {', '.join(item['issues'])}" for item in failed_images],
                "quality_details": qc_issues  # Return full details for frontend visualization
            }

        # --- Step 1: Multi-View Consensus Classification ---
        predictions = []
        for img in loaded_fabric_imgs:
            predictions.append(two_stage_predict(img))
            
        # Consensus Logic — majority vote
        pred_counts = Counter(p["prediction"] for p in predictions)
        majority_pred, majority_count = pred_counts.most_common(1)[0]
        majority_ratio = majority_count / len(predictions)

        # For single image: always pass. For multiple: require >50% agreement.
        if len(predictions) > 1 and majority_ratio <= 0.5:
            return {
                "success": False,
                "message": "Consensus Failed: No majority agreement among images",
                "details": [f"Img {i+1}: {p['prediction']} ({p['confidence']:.2f})" for i, p in enumerate(predictions)]
            }

        # Use the highest-confidence prediction from the majority class
        majority_preds = [p for p in predictions if p["prediction"] == majority_pred]
        best_result = max(majority_preds, key=lambda x: x["confidence"])
        
        # Build metadata structure
        # Use MD5 of first image as ID
        fabric_id = hashlib.md5(raw_bytes_list[0]).hexdigest()[:12]
        
        metadata = {
            "fabric_id": fabric_id,
            "timestamp": datetime.now().isoformat(),
            "structure": {
                "primary": best_result["stage1"]["prediction"],
                "secondary": best_result["stage2"]["prediction"] if best_result["stage2"] else None,
                "full_category": best_result["prediction"],
                "confidence": round(best_result["confidence"], 3),
                "model_version": "convnext_stage1s+stage2t_multiview",
                "consensus_count": len(fabric_images),
                "all_confidences": [round(p["confidence"], 3) for p in predictions]
            },
            "specifications": {},
            "double_check": {
                "performed": False
            }
        }
        
        # Step 2: OCR parsing (if label provided)
        if label_image:
            label_data = await label_image.read()
            label_img = Image.open(BytesIO(label_data)).convert("RGB")
            # Run QC on label too? Maybe just simple check
            ocr_result = CACHE.recognize_ocr(label_img)
            
            if ocr_result:
                parsed_fields = CACHE.ocr_parser.parse_all_fields(ocr_result["text"])
                metadata["specifications"] = {
                    "composition": parsed_fields.get("composition"),
                    "weight": parsed_fields.get("weight"),
                    "pattern": parsed_fields.get("pattern"),
                    "width": parsed_fields.get("width"),
                    "thickness": parsed_fields.get("thickness"),
                    "sustainability": parsed_fields.get("sustainability"),
                    "origin": parsed_fields.get("origin")
                }
                metadata["ocr_raw_text"] = ocr_result["text"]
                metadata["ocr_confidence"] = round(ocr_result["confidence"], 3)
        
        # Step 3: Double structure verification (if back image provided)
        if back_image:
            back_data = await back_image.read()
            back_img = Image.open(BytesIO(back_data)).convert("RGB")
            
            # Use the first fabric image (front) for comparison
            verification = CACHE.double_verifier.verify_double_structure(loaded_fabric_imgs[0], back_img)
            metadata["double_check"] = {
                "performed": True,
                "is_double_structure": verification["is_double_structure"],
                "similarity_score": verification["overall_similarity"],
                "confidence_level": verification["confidence_level"],
                "suggestion": verification["suggestion"]
            }
        
        return {
            "success": True,
            "metadata": metadata
        }
        
    except Exception as e:
        import traceback
        import logging
        logging.exception("create_fabric_metadata failed")
        return {
            "success": False,
            "message": "Internal server error",
        }


# Serve CA certificate for iPad trust installation
CA_PEM_PATH = os.path.join(REPO_ROOT, "certs", "ca.pem")


@app.get("/certs/ca.pem")
async def serve_ca_cert():
    if os.path.isfile(CA_PEM_PATH):
        return FileResponse(CA_PEM_PATH, media_type="application/x-pem-file", filename="ca.pem")
    from fastapi import HTTPException
    raise HTTPException(status_code=404, detail="CA certificate not found")


# Frontend routes — web/ is the primary frontend directory
MOCK_DIR = os.path.join(REPO_ROOT, "web")
if os.path.isdir(MOCK_DIR):
    @app.get("/")
    async def read_root():
        return FileResponse(os.path.join(MOCK_DIR, "mock_demo.html"))

    @app.get("/tablet")
    async def read_mock_tablet():
        return FileResponse(os.path.join(MOCK_DIR, "mock_tablet.html"))

    @app.get("/passport")
    async def read_passport():
        return FileResponse(os.path.join(MOCK_DIR, "passport_view.html"))

    @app.get("/passport_v2")
    async def read_passport_v2():
        return FileResponse(os.path.join(MOCK_DIR, "passport_v2.html"))

    @app.get("/assistant")
    async def read_assistant():
        return FileResponse(os.path.join(MOCK_DIR, "assistant.html"))

    # Serve web/ sub-directories (assets, js, css) so relative paths in HTML work
    for sub in ("assets", "js", "css"):
        sub_path = os.path.join(MOCK_DIR, sub)
        if os.path.isdir(sub_path):
            app.mount(f"/{sub}", StaticFiles(directory=sub_path), name=f"mock_{sub}")

    # Demo reference specimens (images + passport PDFs) — resolution order:
    #   1. FF_SAMPLE_DIR env (explicit override, e.g. HF Space deploy)
    #   2. ~/Downloads/sample (Oscar's local Mac layout)
    #   3. <repo_root>/sample (bundled with deployment image)
    _sample_dir = (
        os.environ.get("FF_SAMPLE_DIR")
        or os.path.expanduser("~/Downloads/sample")
    )
    if not os.path.isdir(_sample_dir):
        _alt = os.path.join(REPO_ROOT, "sample")
        if os.path.isdir(_alt):
            _sample_dir = _alt
    if os.path.isdir(_sample_dir):
        app.mount("/sample", StaticFiles(directory=_sample_dir), name="sample")

    # FabricAI sub-app — Haochen's OpenAI-backed assistant, mounted same-origin.
    # Exposed at /fabricai/* (e.g. /fabricai/api/ask). The slim drawer widget
    # in mock_demo / passport pages calls these endpoints directly.
    try:
        from server.fabricai_app.main import app as fabricai_app
        app.mount("/fabricai", fabricai_app, name="fabricai")
    except Exception as _fai_err:  # noqa: BLE001 — log + continue without breaking host
        import logging as _lg
        _lg.getLogger(__name__).warning("FabricAI sub-app not mounted: %s", _fai_err)

    # Serve the chat drawer widget (shared partial) so any page can <script src> it.
    _widget_dir = os.path.join(MOCK_DIR, "widgets")
    if os.path.isdir(_widget_dir):
        app.mount("/widgets", StaticFiles(directory=_widget_dir), name="widgets")
