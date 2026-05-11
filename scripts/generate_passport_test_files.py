"""Generate realistic test files for passport upload (section 0 — Material/Component, section 3 — Sustainability Certifications).

Outputs:
  /Users/oscar/Downloads/sample/passport_test/material_spec_sheet.pdf
  /Users/oscar/Downloads/sample/passport_test/gots_certificate.docx
"""

from pathlib import Path
from datetime import date

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("/Users/oscar/Downloads/sample/passport_test")
OUT.mkdir(parents=True, exist_ok=True)


# ────────────────────────────────────────────────────────────────────────────
# 1. Material Specification Sheet — PDF
#    Targets DPP Section 0 / "Material and Component List"
#    Fields populated: fibre_composition, yarn specs, weight, width, structure,
#                      shrinkage, durability, colorfastness, supplier.
# ────────────────────────────────────────────────────────────────────────────

def build_pdf():
    path = OUT / "material_spec_sheet.pdf"
    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        topMargin=18 * mm, bottomMargin=18 * mm,
        leftMargin=20 * mm, rightMargin=20 * mm,
        title="Material Specification Sheet",
        author="Hengyuan Textile Technology Co., Ltd.",
    )
    styles = getSampleStyleSheet()
    h_brand = ParagraphStyle('brand', parent=styles['Heading2'],
                             fontSize=10, leading=12, textColor=colors.HexColor('#5b6b4f'),
                             alignment=TA_LEFT, spaceAfter=2)
    h_title = ParagraphStyle('title', parent=styles['Heading1'],
                             fontSize=20, leading=24, textColor=colors.HexColor('#1c2818'),
                             alignment=TA_LEFT, spaceAfter=4)
    h_sub = ParagraphStyle('sub', parent=styles['BodyText'],
                           fontSize=9, leading=12, textColor=colors.HexColor('#666'),
                           alignment=TA_LEFT, spaceAfter=10)
    h_section = ParagraphStyle('section', parent=styles['Heading2'],
                               fontSize=11, leading=14, textColor=colors.HexColor('#28422E'),
                               spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle('body', parent=styles['BodyText'],
                          fontSize=9.5, leading=13, textColor=colors.HexColor('#222'))

    story = []
    story.append(Paragraph("HENGYUAN TEXTILE TECHNOLOGY CO., LTD.", h_brand))
    story.append(Paragraph("Material Specification Sheet", h_title))
    story.append(Paragraph(
        "Document No. HY-MSS-2026-0042 · Issued 2026-04-18 · Buyer: FabricFlow Demo Account",
        h_sub))

    # 1. Product Identity
    story.append(Paragraph("1. Product Identity", h_section))
    rows = [
        ["Fabric Name", "Heritage Cotton Twill 168"],
        ["Internal SKU", "HY-CTW-168-NAT"],
        ["Passport ID / Lot", "FF-001 / Batch 2026-04-Q2-018"],
        ["Country of Origin", "China (Guangdong, Shaoguan)"],
        ["Category", "Woven · Twill (Z-twill, 3/1)"],
        ["Updated", "2026-04-18"],
    ]
    story.append(_kv_table(rows))

    # 2. Fibre Composition
    story.append(Paragraph("2. Fibre Composition", h_section))
    fib = [
        ["Fibre", "Percentage", "Origin", "Note"],
        ["GOTS-certified Organic Cotton", "62 %", "Xinjiang, China", "Long-staple, GOTS scope cert. CN-BCS-090421"],
        ["Recycled Linen (post-industrial)", "38 %", "Heilongjiang, China", "Mechanically reclaimed, GRS-certified"],
    ]
    story.append(_grid_table(fib, col_widths=[58 * mm, 22 * mm, 38 * mm, 52 * mm]))

    # 3. Yarn Specifications
    story.append(Paragraph("3. Yarn Specifications", h_section))
    rows = [
        ["Yarn Count", "Nm 40/2 (warp) · Nm 36/2 (weft)"],
        ["Yarn Twist", "Z-twist 720 tpm (warp) · S-twist 680 tpm (weft)"],
        ["Spinning Method", "Compact ring-spun"],
        ["Fibre Length / Micronaire", "30.5 mm staple · 4.1 micronaire"],
        ["Trash Content", "0.8 %"],
    ]
    story.append(_kv_table(rows))

    # 4. Physical Specifications
    story.append(Paragraph("4. Physical Specifications", h_section))
    rows = [
        ["Fabric Weight", "168 ± 4 g/m² (ISO 3801)"],
        ["Fabric Width", "148 cm (cuttable 144 cm)"],
        ["Fabric Structure", "3/1 Z Twill, 78 × 64 ends/inch"],
        ["Colour State", "Natural greige (un-dyed)"],
    ]
    story.append(_kv_table(rows))

    # 5. Dimensional Stability
    story.append(Paragraph("5. Dimensional Stability — ISO 6330", h_section))
    rows = [
        ["Shrinkage (Warp)", "−2.4 %  (5 wash 40°C tumble dry)"],
        ["Shrinkage (Weft)", "−1.6 %"],
    ]
    story.append(_kv_table(rows))

    # 6. Durability
    story.append(Paragraph("6. Durability Tests", h_section))
    dur = [
        ["Test", "Result", "Method"],
        ["Pilling Resistance", "Grade 4 (5 = best)", "ISO 12945-2 Martindale"],
        ["Abrasion Resistance", "32,000 cycles (no breakdown)", "ISO 12947-2"],
        ["Tensile Strength (Warp)", "812 N", "ISO 13934-1 strip"],
        ["Tear Strength", "48 N", "ISO 13937-1 ballistic"],
    ]
    story.append(_grid_table(dur, col_widths=[55 * mm, 55 * mm, 60 * mm]))

    # 7. Colourfastness — ISO 105
    story.append(Paragraph("7. Colour Fastness — ISO 105", h_section))
    cf = [
        ["Property", "Grade", "Note"],
        ["Washing", "4–5", "ISO 105-C06 A1S"],
        ["Rubbing (Dry)", "4", "ISO 105-X12"],
        ["Rubbing (Wet)", "3–4", "ISO 105-X12"],
        ["Light", "5", "ISO 105-B02"],
    ]
    story.append(_grid_table(cf, col_widths=[55 * mm, 30 * mm, 85 * mm]))

    # 8. Supplier
    story.append(Paragraph("8. Supplier Information", h_section))
    rows = [
        ["Supplier Name", "Guangdong Hengyuan Textile Technology Co., Ltd."],
        ["Location", "Shaoguan Industrial Park, Guangdong, China"],
        ["Tier", "Tier 2 — fabric mill"],
        ["Annual Capacity", "12.4 million metres"],
        ["Workforce", "428 staff"],
        ["Last Audit", "2026-02-11 · Pass (SMETA 4-Pillar)"],
        ["On-Time Delivery", "96.4 %"],
        ["Partnership Years", "7 years"],
    ]
    story.append(_kv_table(rows))

    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "<i>This document is issued for FabricFlow demo evaluation. "
        "All test results derive from ISO-accredited laboratory reports on file. "
        "Refer to certification annexes for GOTS, GRS, and OEKO-TEX scope.</i>",
        body))

    doc.build(story)
    print(f"wrote {path}")


def _kv_table(rows):
    t = Table(rows, colWidths=[55 * mm, 115 * mm])
    t.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 9.2),
        ('FONT', (0, 0), (0, -1), 'Helvetica-Bold', 9.2),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#3b4a31')),
        ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor('#222')),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#f7f7f0'), colors.white]),
        ('LINEABOVE', (0, 0), (-1, 0), 0.5, colors.HexColor('#cbd2bf')),
        ('LINEBELOW', (0, -1), (-1, -1), 0.5, colors.HexColor('#cbd2bf')),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    return t


def _grid_table(rows, col_widths=None):
    t = Table(rows, colWidths=col_widths)
    t.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 9.2),
        ('FONT', (0, 1), (-1, -1), 'Helvetica', 9.2),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#28422E')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#222')),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#cbd2bf')),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    return t


# ────────────────────────────────────────────────────────────────────────────
# 2. GOTS / Sustainability Certificate — DOCX
#    Targets DPP Section 3 / "Sustainability Certifications"
#    Fields: GOTS organic content, GRS recycled content, supplier audit,
#            ZDHC level, OEKO-TEX standard 100, environmental impact context.
# ────────────────────────────────────────────────────────────────────────────

def build_docx():
    path = OUT / "gots_certificate.docx"
    d = Document()

    # Page setup
    section = d.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Header
    h0 = d.add_paragraph()
    h0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h0.add_run("CONTROL UNION CERTIFICATIONS B.V.")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x5b, 0x6b, 0x4f)

    title = d.add_paragraph()
    rt = title.add_run("Scope Certificate — GOTS · GRS")
    rt.bold = True
    rt.font.size = Pt(20)
    rt.font.color.rgb = RGBColor(0x1c, 0x28, 0x18)

    sub = d.add_paragraph()
    rs = sub.add_run("Certificate No. CU 2026-CN-840271 · Valid 2026-04-01 → 2027-03-31 · Issued The Hague, Netherlands")
    rs.font.size = Pt(9)
    rs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _hr(d)

    # 1. Certificate Holder
    _section(d, "1. Certificate Holder")
    _kv(d, [
        ("Company", "Guangdong Hengyuan Textile Technology Co., Ltd."),
        ("Address", "Shaoguan Industrial Park, Guangdong Province, China"),
        ("Registration", "GD-91440205-MA7K8X12"),
        ("Tier", "Tier 2 — fabric weaver and finisher"),
    ])

    # 2. Standards & Scope
    _section(d, "2. Standards & Certified Scope")
    _kv(d, [
        ("GOTS Version", "GOTS 7.0 (2025 revision)"),
        ("GOTS Scope", "Weaving, dyeing, finishing of organic cotton woven fabrics"),
        ("Organic Content (declared)", "62 % certified organic cotton (in-conversion ≤ 10 %)"),
        ("GRS Version", "GRS 4.0"),
        ("GRS Scope", "Mechanical recycling and use of recycled linen fibre"),
        ("Recycled Content (declared)", "38 % post-industrial recycled linen"),
        ("OEKO-TEX STANDARD 100", "Certificate 22.HCN.94217 · Product Class I (skin-contact)"),
    ])

    # 3. Material declared on certificate
    _section(d, "3. Materials Covered")
    _table(d, [
        ["Fibre", "Composition %", "Status"],
        ["Cotton (organic, GOTS)", "62 %", "Certified — in scope"],
        ["Linen (recycled, GRS)", "38 %", "Certified — in scope"],
    ])

    # 4. Chemical Inputs Compliance
    _section(d, "4. Chemical Inputs — ZDHC Conformity")
    _kv(d, [
        ("MRSL Conformance Level", "ZDHC MRSL v3.1 — Level 3 (highest)"),
        ("Dyestuff Class", "Reactive bi-functional dyestuffs (GOTS-approved)"),
        ("Pretreatment", "Enzymatic desize / scour, no chlorine bleach"),
        ("Softener", "Silicone-free, biodegradable cationic softener"),
        ("Heat Setting", "Saturated steam, 130 °C / 90 s — no PFAS, no formaldehyde donor"),
    ])

    # 5. Audit & Compliance
    _section(d, "5. Audit & Compliance Summary")
    _kv(d, [
        ("Last Audit Date", "2026-02-11"),
        ("Audit Standard", "SMETA 4-Pillar (SEDEX)"),
        ("Result", "Pass — 0 non-conformances, 2 observations closed within 14 days"),
        ("REACH Status", "Compliant — Annex XIV / XVII screened, no SVHC > 0.1 %"),
        ("EU CoC Status", "Conforms with Regulation (EU) 1007/2011 on textile labelling"),
    ])

    # 6. Environmental Impact (declared by holder, audited)
    _section(d, "6. Declared Environmental Impact (per kg finished fabric)")
    _table(d, [
        ["Indicator", "Declared Value", "Industry Benchmark"],
        ["Water consumption", "62 L/kg", "120 L/kg (cotton woven avg.)"],
        ["Carbon footprint", "4.8 kg CO₂e/kg", "8.1 kg CO₂e/kg"],
        ["Energy demand", "38 MJ/kg", "55 MJ/kg"],
        ["Wastewater recovery", "74 %", "30 %"],
        ["Renewable electricity share", "41 % (rooftop PV + grid PPA)", "≈ 12 %"],
    ])

    # 7. Validity Statement
    _section(d, "7. Validity Statement")
    p = d.add_paragraph()
    p.add_run(
        "This Scope Certificate confirms that the named operator's processing units, "
        "as listed under §1, were inspected on the audit date stated in §5 and found in compliance "
        "with the GOTS 7.0, GRS 4.0, and OEKO-TEX STANDARD 100 schemes. Validity is subject to "
        "annual surveillance audits and the operator's continued compliance with chemical inputs, "
        "social responsibility, and traceability requirements."
    ).font.size = Pt(9.5)

    # Footer
    d.add_paragraph()
    f = d.add_paragraph()
    rf = f.add_run("Issued for FabricFlow demonstration purposes. Verifiable at controlunion.com/certificate/CU-2026-CN-840271.")
    rf.italic = True
    rf.font.size = Pt(8.5)
    rf.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    d.save(path)
    print(f"wrote {path}")


def _hr(d):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("─" * 70)
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xcb, 0xd2, 0xbf)


def _section(d, title):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x28, 0x42, 0x2e)


def _kv(d, rows):
    t = d.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c0 = t.rows[i].cells[0]
        c1 = t.rows[i].cells[1]
        c0.width = Cm(5.5)
        c1.width = Cm(11.5)
        for c, txt, bold in ((c0, k, True), (c1, v, False)):
            cp = c.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            run = cp.add_run(txt)
            run.font.size = Pt(9.5)
            run.bold = bold
            if bold:
                run.font.color.rgb = RGBColor(0x3b, 0x4a, 0x31)


def _table(d, rows):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.rows[i].cells[j]
            cp = c.paragraphs[0]
            run = cp.add_run(cell)
            run.font.size = Pt(9.5)
            run.bold = (i == 0)
            if i == 0:
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                _shade(c, "28422E")


def _shade(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


# ────────────────────────────────────────────────────────────────────────────
# 3. Purchase Contract — PDF (Section 1 / Procurement & Commercial Documents)
# ────────────────────────────────────────────────────────────────────────────

def build_contract_pdf():
    path = OUT / "purchase_contract.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=20 * mm, rightMargin=20 * mm,
                            title="Fabric Purchase Contract", author="FabricFlow Demo")
    s = getSampleStyleSheet()
    h_brand = ParagraphStyle('b', parent=s['Heading2'], fontSize=10, leading=12,
                             textColor=colors.HexColor('#5b6b4f'), spaceAfter=2)
    h_title = ParagraphStyle('t', parent=s['Heading1'], fontSize=20, leading=24,
                             textColor=colors.HexColor('#1c2818'), spaceAfter=4)
    h_sub = ParagraphStyle('su', parent=s['BodyText'], fontSize=9, leading=12,
                           textColor=colors.HexColor('#666'), spaceAfter=10)
    h_section = ParagraphStyle('sec', parent=s['Heading2'], fontSize=11,
                               textColor=colors.HexColor('#28422E'), spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle('bd', parent=s['BodyText'], fontSize=9.5, leading=13,
                          textColor=colors.HexColor('#222'))
    story = [
        Paragraph("FABRICFLOW DEMO ACCOUNT — BRAND BUYER", h_brand),
        Paragraph("Fabric Purchase Contract", h_title),
        Paragraph("Contract No. FF-PO-2026-0312 · Issued 2026-04-22 · Incoterm CIF Rotterdam", h_sub),
        Paragraph("1. Parties", h_section),
        _kv_table([
            ["Buyer", "FabricFlow Demo Account, London EC1V 9BX, UK"],
            ["Seller", "Guangdong Hengyuan Textile Technology Co., Ltd."],
            ["Country of Origin", "China (Guangdong)"],
            ["Bank Account", "Bank of China, Shaoguan Branch · Account 4382 1098 7263"],
        ]),
        Paragraph("2. Goods", h_section),
        _grid_table([
            ["Item", "Description", "Quantity", "Unit Price", "Total"],
            ["1", "Heritage Cotton Twill 168 (HY-CTW-168-NAT)", "4,800 m", "USD 14.20", "USD 68,160.00"],
            ["2", "Quality test sampling allowance", "20 m", "Free of charge", "—"],
        ], col_widths=[16 * mm, 70 * mm, 26 * mm, 28 * mm, 30 * mm]),
        Paragraph("3. Specifications", h_section),
        _kv_table([
            ["Composition", "62% GOTS Organic Cotton / 38% Recycled Linen"],
            ["Weight", "168 ± 4 g/m² (ISO 3801)"],
            ["Width", "148 cm cuttable"],
            ["Construction", "3/1 Z Twill, 78 × 64 ends/inch"],
            ["Lot Reference", "FF-001 / Batch 2026-04-Q2-018"],
        ]),
        Paragraph("4. Commercial Terms", h_section),
        _kv_table([
            ["Incoterm", "CIF Rotterdam (Incoterms 2020)"],
            ["Payment Terms", "30 % advance, 70 % against B/L copy via T/T"],
            ["Delivery Window", "Latest shipment 2026-06-10"],
            ["Lead Time", "35–45 days from deposit confirmation"],
            ["Currency", "USD"],
            ["Total Contract Value", "USD 68,160.00"],
        ]),
        Paragraph("5. Quality & Compliance", h_section),
        Paragraph(
            "Goods shall conform to the specification annex and to GOTS 7.0, GRS 4.0, and OEKO-TEX STANDARD 100 "
            "Product Class I. Seller warrants that no SVHC substances above 0.1 % m/m are present. Inspection rights "
            "reserved on production line and at port of loading. Non-conformity claims may be raised within 14 days "
            "of receipt at Rotterdam.",
            body),
        Paragraph("6. Signatures", h_section),
        _kv_table([
            ["Buyer Signature", "_______________________  ·  Date 2026-04-22"],
            ["Seller Signature", "_______________________  ·  Date 2026-04-22"],
        ]),
    ]
    doc.build(story)
    print(f"wrote {path}")


# ────────────────────────────────────────────────────────────────────────────
# 4. Bill of Lading — PDF (Section 1 / Logistics & Shipping)
# ────────────────────────────────────────────────────────────────────────────

def build_bol_pdf():
    path = OUT / "bill_of_lading.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=20 * mm, rightMargin=20 * mm)
    s = getSampleStyleSheet()
    h_brand = ParagraphStyle('b', parent=s['Heading2'], fontSize=10,
                             textColor=colors.HexColor('#5b6b4f'), spaceAfter=2)
    h_title = ParagraphStyle('t', parent=s['Heading1'], fontSize=20, leading=24,
                             textColor=colors.HexColor('#1c2818'), spaceAfter=4)
    h_sub = ParagraphStyle('su', parent=s['BodyText'], fontSize=9, leading=12,
                           textColor=colors.HexColor('#666'), spaceAfter=10)
    h_section = ParagraphStyle('sec', parent=s['Heading2'], fontSize=11,
                               textColor=colors.HexColor('#28422E'), spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle('bd', parent=s['BodyText'], fontSize=9.5, leading=13,
                          textColor=colors.HexColor('#222'))
    story = [
        Paragraph("MAERSK LINE A/S", h_brand),
        Paragraph("Bill of Lading", h_title),
        Paragraph("B/L No. MAEU-983241756 · Issued 2026-05-12 Yantian, China · Original 3 of 3", h_sub),
        Paragraph("Shipper / Consignee", h_section),
        _kv_table([
            ["Shipper", "Guangdong Hengyuan Textile Technology Co., Ltd. · Shaoguan, China"],
            ["Consignee", "FabricFlow Demo Account · London EC1V 9BX, UK"],
            ["Notify Party", "DHL Forwarding (Rotterdam) — ref FF-PO-2026-0312"],
        ]),
        Paragraph("Routing", h_section),
        _kv_table([
            ["Vessel / Voyage", "MAERSK ENSHADEN · 2615W"],
            ["Port of Loading", "Yantian, Shenzhen (CNYTN)"],
            ["Port of Discharge", "Rotterdam, Netherlands (NLRTM)"],
            ["Place of Receipt", "Shaoguan ICD"],
            ["Place of Delivery", "Rotterdam (door)"],
            ["Departure", "2026-05-12"],
            ["ETA", "2026-06-08 (transit ≈ 27 days)"],
        ]),
        Paragraph("Cargo", h_section),
        _grid_table([
            ["Marks & Nos.", "Pkg & Description", "Gross kg", "Measurement", "HS Code"],
            ["FF/PO-2026-0312", "12 × Wooden crates · Cotton Twill Greige Fabric", "1,140 kg", "5.6 m³", "5208.39"],
        ], col_widths=[34 * mm, 70 * mm, 22 * mm, 26 * mm, 22 * mm]),
        Paragraph("Container Details", h_section),
        _kv_table([
            ["Container No.", "MSKU 8821037 (40' HC)"],
            ["Seal No.", "ML-7732499"],
            ["Freight Terms", "Freight Prepaid (CIF Rotterdam)"],
            ["Insurance", "Marine cargo · ICC(A) Lloyd's, policy GB-2026-440021"],
        ]),
        Paragraph("Declarations", h_section),
        Paragraph(
            "Country of Origin: China (Guangdong). Goods clean on board, in apparent good order and condition unless "
            "otherwise noted. Subject to terms and conditions of the carrier's Bill of Lading. Cargo declared free of "
            "hazardous goods (IMO not applicable).",
            body),
    ]
    doc.build(story)
    print(f"wrote {path}")


# ────────────────────────────────────────────────────────────────────────────
# 5. OEKO-TEX Certificate — PDF (Section 3 / Sustainability Certifications)
# ────────────────────────────────────────────────────────────────────────────

def build_oekotex_pdf():
    path = OUT / "oeko_tex_certificate.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=20 * mm, rightMargin=20 * mm)
    sty = getSampleStyleSheet()
    h_brand = ParagraphStyle('b', parent=sty['Heading2'], fontSize=10,
                             textColor=colors.HexColor('#5b6b4f'), spaceAfter=2)
    h_title = ParagraphStyle('t', parent=sty['Heading1'], fontSize=20, leading=24,
                             textColor=colors.HexColor('#1c2818'), spaceAfter=4)
    h_sub = ParagraphStyle('su', parent=sty['BodyText'], fontSize=9,
                           textColor=colors.HexColor('#666'), spaceAfter=10)
    h_section = ParagraphStyle('sec', parent=sty['Heading2'], fontSize=11,
                               textColor=colors.HexColor('#28422E'), spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle('bd', parent=sty['BodyText'], fontSize=9.5, leading=13,
                          textColor=colors.HexColor('#222'))
    story = [
        Paragraph("OEKO-TEX SERVICE GMBH · TESTEX AG", h_brand),
        Paragraph("STANDARD 100 by OEKO-TEX®", h_title),
        Paragraph("Certificate No. 22.HCN.94217 · Valid 2026-04-01 → 2027-03-31 · Test Institute Hong Kong", h_sub),
        Paragraph("1. Certificate Holder", h_section),
        _kv_table([
            ["Company", "Guangdong Hengyuan Textile Technology Co., Ltd."],
            ["Address", "Shaoguan Industrial Park, Guangdong, China"],
            ["Customer ID", "OK-CN-22841"],
        ]),
        Paragraph("2. Tested Article", h_section),
        _kv_table([
            ["Article", "Heritage Cotton Twill 168 (HY-CTW-168-NAT)"],
            ["Composition", "62 % Organic Cotton / 38 % Recycled Linen"],
            ["Construction", "Woven, plain twill 3/1, 168 g/m²"],
            ["Article Colours", "Greige (un-dyed), white-dyed, navy reactive"],
            ["Product Class", "I — Articles for babies and small children (most stringent)"],
        ]),
        Paragraph("3. Test Criteria — Limit Values Verified", h_section),
        _grid_table([
            ["Parameter", "Limit (Class I)", "Result"],
            ["pH value", "4.0 – 7.5", "6.4 ✓"],
            ["Formaldehyde (free + released)", "≤ 16 mg/kg", "Not detected ✓"],
            ["Heavy metals (Lead, Cadmium, Mercury, …)", "ICP-MS", "Below limits ✓"],
            ["Pesticides (sum)", "≤ 0.5 mg/kg", "Not detected ✓"],
            ["Pentachlorophenol (PCP)", "≤ 0.05 mg/kg", "Not detected ✓"],
            ["Allergenic disperse dyes", "Not detectable", "Not detected ✓"],
            ["Carcinogenic dyes (incl. azo)", "Not detectable", "Not detected ✓"],
            ["VOC emissions", "GC-MS", "Below limits ✓"],
            ["Odour test", "Pass / Fail", "Pass ✓"],
            ["Colour fastness — saliva and perspiration", "≥ 4", "Grade 4–5 ✓"],
        ], col_widths=[80 * mm, 50 * mm, 40 * mm]),
        Paragraph("4. Validity Statement", h_section),
        Paragraph(
            "The above article complies with the human-ecological requirements of the STANDARD 100 by OEKO-TEX® "
            "currently established in Annex 4 (Edition 04/2026) for Product Class I. This certificate is valid for "
            "12 months. The certificate holder shall comply with the conformity-control plan in line with §7 of the "
            "STANDARD 100 by OEKO-TEX® regulations.",
            body),
        Paragraph("Issued 2026-04-01, Hong Kong · Authorised Signatory: Dr Wai-Lun Cheung, Test Director.", body),
    ]
    doc.build(story)
    print(f"wrote {path}")


# ────────────────────────────────────────────────────────────────────────────
# 6. REACH Compliance Declaration — DOCX (Section 3 / Regulatory)
# ────────────────────────────────────────────────────────────────────────────

def build_reach_docx():
    path = OUT / "reach_compliance_declaration.docx"
    d = Document()
    sec = d.sections[0]
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

    h0 = d.add_paragraph(); r = h0.add_run("HENGYUAN TEXTILE TECHNOLOGY CO., LTD."); r.bold = True
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x5b, 0x6b, 0x4f)
    t = d.add_paragraph(); rt = t.add_run("REACH Compliance Declaration"); rt.bold = True
    rt.font.size = Pt(20); rt.font.color.rgb = RGBColor(0x1c, 0x28, 0x18)
    su = d.add_paragraph(); rs = su.add_run(
        "Declaration ID HY-REACH-2026-088 · Issued 2026-04-15 · Pertaining to Regulation (EC) No 1907/2006")
    rs.font.size = Pt(9); rs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _hr(d)

    _section(d, "1. Declarant")
    _kv(d, [
        ("Company", "Guangdong Hengyuan Textile Technology Co., Ltd."),
        ("Address", "Shaoguan Industrial Park, Guangdong, China"),
        ("Authorised EU Representative", "Textile Compliance B.V., Amsterdam"),
    ])
    _section(d, "2. Article Declared")
    _kv(d, [
        ("Article Name", "Heritage Cotton Twill 168"),
        ("SKU", "HY-CTW-168-NAT"),
        ("Composition", "62 % Organic Cotton / 38 % Recycled Linen"),
        ("HS Code", "5208.39"),
    ])
    _section(d, "3. SVHC Screening — Annex XIV / Annex XVII")
    _table(d, [
        ["Substance Group", "Threshold", "Result"],
        ["Phthalates (DEHP, BBP, DBP, DIBP)", "≤ 0.1 % m/m", "Not detected"],
        ["Polycyclic aromatic hydrocarbons (PAHs)", "≤ 1 mg/kg", "Not detected"],
        ["Nonylphenol ethoxylates (NPEO)", "≤ 0.01 % m/m", "Not detected"],
        ["Per- and polyfluoroalkyl substances (PFAS)", "Detection limit", "Not detected"],
        ["Azo dyes releasing aromatic amines", "Not detectable", "Not detected"],
        ["Organotin compounds", "Below limit", "Not detected"],
        ["Chromium VI (leather only — N/A)", "n/a", "Not applicable"],
    ])
    _section(d, "4. Statement of Conformity")
    p = d.add_paragraph(); p.add_run(
        "Hengyuan Textile Technology Co., Ltd. hereby declares that the above article, in the form supplied to FabricFlow "
        "under contract FF-PO-2026-0312, complies with Regulation (EC) No 1907/2006 (REACH), the latest Candidate List of "
        "Substances of Very High Concern (SVHCs) issued by ECHA, and Annex XVII restrictions applicable to textile articles. "
        "No SVHC is intentionally added above 0.1 % m/m of the article."
    ).font.size = Pt(9.5)
    _section(d, "5. Supporting Test Reports")
    _kv(d, [
        ("Lab", "Bureau Veritas Consumer Products Services, Hong Kong"),
        ("Report ID", "BV-2026-HK-441207"),
        ("Test Methods", "EN 14582, EN ISO 17226, EN 16711"),
        ("Test Date", "2026-03-29"),
    ])
    d.add_paragraph()
    f = d.add_paragraph(); rf = f.add_run(
        "Issued for FabricFlow demonstration purposes. Authorised signatory: Mr Yong Liu, Quality Director, Hengyuan Textile.")
    rf.italic = True; rf.font.size = Pt(8.5); rf.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    d.save(path); print(f"wrote {path}")


# ────────────────────────────────────────────────────────────────────────────
# 7. LCA Report — PDF (Section 4 / Lifecycle & Environmental Footprint)
# ────────────────────────────────────────────────────────────────────────────

def build_lca_pdf():
    path = OUT / "lca_report.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=20 * mm, rightMargin=20 * mm)
    s = getSampleStyleSheet()
    h_brand = ParagraphStyle('b', parent=s['Heading2'], fontSize=10,
                             textColor=colors.HexColor('#5b6b4f'), spaceAfter=2)
    h_title = ParagraphStyle('t', parent=s['Heading1'], fontSize=20, leading=24,
                             textColor=colors.HexColor('#1c2818'), spaceAfter=4)
    h_sub = ParagraphStyle('su', parent=s['BodyText'], fontSize=9,
                           textColor=colors.HexColor('#666'), spaceAfter=10)
    h_section = ParagraphStyle('sec', parent=s['Heading2'], fontSize=11,
                               textColor=colors.HexColor('#28422E'), spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle('bd', parent=s['BodyText'], fontSize=9.5, leading=13,
                          textColor=colors.HexColor('#222'))
    story = [
        Paragraph("QUANTIS / EARTHSTER · LCA SERVICE", h_brand),
        Paragraph("Cradle-to-Gate Life Cycle Assessment", h_title),
        Paragraph("Report ID QTS-LCA-2026-1187 · Issued 2026-04-09 · Functional unit 1 kg finished fabric", h_sub),
        Paragraph("1. Goal & Scope", h_section),
        _kv_table([
            ["Article", "Heritage Cotton Twill 168 (HY-CTW-168-NAT)"],
            ["Functional Unit", "1 kg finished fabric, ready for cut and sew"],
            ["System Boundary", "Cradle-to-gate (raw fibre → packed at mill)"],
            ["Methodology", "ISO 14040 / 14044, PEFCR Apparel & Footwear v1.3"],
            ["Database", "ecoinvent 3.10, Higg MSI 3.7"],
        ]),
        Paragraph("2. Headline Indicators", h_section),
        _grid_table([
            ["Indicator", "Result", "Industry Benchmark", "Δ vs benchmark"],
            ["Carbon footprint (GWP100)", "4.8 kg CO₂e", "8.1 kg CO₂e", "−41 %"],
            ["Water consumption (blue water)", "62 L", "120 L", "−48 %"],
            ["Energy demand (CED)", "38 MJ", "55 MJ", "−31 %"],
            ["Land use", "1.9 m²·a", "3.2 m²·a", "−41 %"],
            ["Eutrophication (freshwater)", "0.0021 kg P-eq", "0.0044", "−52 %"],
            ["Microplastic shedding (cotton-dominant)", "Low", "Medium-High", "Better"],
        ], col_widths=[60 * mm, 32 * mm, 38 * mm, 30 * mm]),
        Paragraph("3. Life Cycle Stage Contribution (% of GWP)", h_section),
        _grid_table([
            ["Stage", "GWP share"],
            ["Fibre cultivation (organic cotton + recycled linen reclaim)", "32 %"],
            ["Yarn spinning", "14 %"],
            ["Weaving", "11 %"],
            ["Dyeing & finishing", "29 %"],
            ["Logistics (mill ↔ port)", "9 %"],
            ["Packaging & ancillary", "5 %"],
        ], col_widths=[110 * mm, 50 * mm]),
        Paragraph("4. Critical Assumptions & Data Quality", h_section),
        Paragraph(
            "Energy mix: 41 % renewable (rooftop PV + grid PPA), 49 % grid (China South Power Grid 2025 mix), 10 % "
            "gas. Wastewater treatment recovery at facility = 74 %. ZDHC MRSL Level 3 conformance verified. "
            "Recycled linen impact allocated 50/50 with primary linen production (cut-off allocation per ISO 14044). "
            "Data quality pedigree score = 1.6 (Weidema scale) — high quality.",
            body),
        Paragraph("5. Verification & Use", h_section),
        Paragraph(
            "This LCA report has been independently verified by SCS Global Services on 2026-04-05 (verification "
            "statement SCS-V-22919). Results are intended for internal product stewardship and B2B disclosure under "
            "the EU ESPR Delegated Act anticipated 2027. Comparative claims against competing products require a "
            "separate critical-review-panel report.",
            body),
    ]
    doc.build(story)
    print(f"wrote {path}")


# ────────────────────────────────────────────────────────────────────────────
# 8. Care Label — image (Section 2 / Care and Use Instructions)
# ────────────────────────────────────────────────────────────────────────────

def build_care_label():
    import os as _os
    from PIL import Image as PILImage, ImageDraw, ImageFont
    W, H = 900, 1300
    img = PILImage.new("RGB", (W, H), (252, 248, 240))
    d = ImageDraw.Draw(img)
    paths = ["/System/Library/Fonts/Helvetica.ttc",
             "/System/Library/Fonts/Supplemental/Arial.ttf"]
    def font(sz, bold=False):
        for pth in paths:
            if _os.path.exists(pth): return ImageFont.truetype(pth, sz)
        return ImageFont.load_default()

    d.rectangle([0, 0, W, 90], fill=(40, 66, 46))
    d.text((30, 28), "HENGYUAN TEXTILE — CARE LABEL", fill=(255, 255, 255), font=font(24, True))
    d.text((30, 130), "Article: Heritage Cotton Twill 168 (HY-CTW-168-NAT)", fill=(40, 40, 40), font=font(20))
    d.text((30, 170), "Composition: 62% Organic Cotton / 38% Recycled Linen", fill=(40, 40, 40), font=font(20))
    d.text((30, 210), "Lot: FF-001 · 2026-04-Q2-018", fill=(40, 40, 40), font=font(20))
    d.line([30, 260, W - 30, 260], fill=(200, 200, 180), width=2)

    care_items = [
        ("Washing", "Machine wash cold ≤ 30 °C, gentle cycle, similar colours."),
        ("Bleaching", "Do not bleach. Avoid optical brighteners."),
        ("Drying", "Tumble dry low or line dry in shade. Reshape whilst damp."),
        ("Ironing", "Iron warm ≤ 150 °C. Steam permissible from reverse."),
        ("Dry Cleaning", "Professional dry clean P (perchloroethylene) optional."),
        ("First Wash", "Expect 2–3 % relaxation shrinkage; pre-shrunk."),
        ("Storage", "Store flat, away from direct sunlight."),
    ]
    y = 290
    for k, v in care_items:
        d.text((30, y), k + ":", fill=(40, 66, 46), font=font(20, True))
        d.text((230, y), v, fill=(20, 20, 20), font=font(20))
        y += 60

    d.line([30, y + 10, W - 30, y + 10], fill=(200, 200, 180), width=2)
    d.text((30, y + 30), "Sustainability: GOTS Organic, OEKO-TEX 100, GRS",
           fill=(40, 66, 46), font=font(20, True))
    d.text((30, y + 70), "Origin: Made in China (Guangdong)",
           fill=(40, 66, 46), font=font(20, True))
    d.text((30, y + 110), "End-of-life: Mechanically recyclable (mono-fibre garment recommended)",
           fill=(40, 66, 46), font=font(20))

    d.rectangle([0, H - 70, W, H], fill=(40, 66, 46))
    d.text((30, H - 50), "FabricFlow Demo · Care & Circularity Label",
           fill=(255, 255, 255), font=font(18))
    p = OUT / "care_label.png"
    img.save(p)
    print(f"wrote {p}")


if __name__ == "__main__":
    build_pdf()
    build_docx()
    build_contract_pdf()
    build_bol_pdf()
    build_oekotex_pdf()
    build_reach_docx()
    build_lca_pdf()
    build_care_label()
